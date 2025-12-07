import streamlit as st
import pandas as pd
import re
import datetime
from io import BytesIO
from docx import Document
import json
import os
from PIL import Image  

# ============================================
# PAGE CONFIGURATION
# ============================================
def _get_page_icon():
    # מפעיל חיפוש גם מקומי וגם ב-/mnt/data למקרה של הרצה בסביבה אחרת
    for path in ["Capture.PNG", "/mnt/data/Capture.PNG"]:
        if os.path.exists(path):
            return Image.open(path)
    return "💼"  # fallback אם הקובץ לא נמצא

st.set_page_config(
    page_title="Aiolos Financial Tools",
    page_icon=_get_page_icon(),
    layout="wide",
    initial_sidebar_state="collapsed"
)

# ============================================
# CUSTOM CSS - CLEAN & MODERN DESIGN
# ============================================
st.markdown("""
<style>
    /* Import modern font */
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap');
    
    /* Reset and base styles */
    * {
        font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif;
    }
    
    /* Main container */
    .main {
        padding: 2rem;
        max-width: 1400px;
        margin: 0 auto;
    }
    
    /* Hide Streamlit branding */
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    header {visibility: hidden;}
    
    /* Custom header with logo */
    .app-header {
        background: linear-gradient(135deg, #1e3c72 0%, #2a5298 100%);
        color: white;
        padding: 1.5rem 2rem;
        border-radius: 12px;
        margin-bottom: 2rem;
        box-shadow: 0 10px 30px rgba(0,0,0,0.1);
        display: flex;
        align-items: center;
        gap: 2rem;
    }
    
    .logo-container {
        flex-shrink: 0;
    }
    
    .logo-container img {
        width: 80px;
        height: 80px;
        border-radius: 50%;
        border: 3px solid rgba(255,255,255,0.3);
        box-shadow: 0 4px 15px rgba(0,0,0,0.2);
    }
    
    .header-text {
        flex-grow: 1;
    }
    
    .app-title {
        font-size: 2rem;
        font-weight: 700;
        margin: 0;
    }
    
    .app-subtitle {
        font-size: 1rem;
        opacity: 0.9;
        margin-top: 0.5rem;
    }
    
    /* Navigation tabs */
    .stTabs [data-baseweb="tab-list"] {
        gap: 1rem;
        background: white;
        padding: 0.5rem;
        border-radius: 10px;
        box-shadow: 0 2px 10px rgba(0,0,0,0.05);
    }
    
    .stTabs [data-baseweb="tab"] {
        height: 50px;
        padding: 0 1.5rem;
        background: transparent;
        border-radius: 8px;
        font-weight: 500;
        color: #666;
    }
    
    .stTabs [data-baseweb="tab"]:hover {
        background: #f8f9fa;
    }
    
    .stTabs [aria-selected="true"] {
        background: linear-gradient(135deg, #1e3c72 0%, #2a5298 100%);
        color: white;
    }
    
    /* Cards */
    .info-card {
        background: white;
        padding: 1.5rem;
        border-radius: 12px;
        box-shadow: 0 2px 10px rgba(0,0,0,0.05);
        border: 1px solid #f0f0f0;
        margin-bottom: 1.5rem;
    }
    
    .info-card h3 {
        margin-top: 0;
        color: #333;
        font-size: 1.2rem;
        font-weight: 600;
    }
    
    /* Template preview */
    .template-preview {
        background: white;
        padding: 1rem;
        border-radius: 12px;
        box-shadow: 0 2px 10px rgba(0,0,0,0.05);
        border: 2px solid #e0e0e0;
        text-align: center;
    }
    
    .template-preview img {
        max-width: 100%;
        border-radius: 8px;
    }
    
    /* Metrics */
    .metric-container {
        display: flex;
        gap: 1.5rem;
        margin: 2rem 0;
    }
    
    .metric-box {
        flex: 1;
        background: white;
        padding: 1.5rem;
        border-radius: 12px;
        border: 1px solid #f0f0f0;
        text-align: center;
        transition: transform 0.2s;
    }
    
    .metric-box:hover {
        transform: translateY(-2px);
        box-shadow: 0 5px 20px rgba(0,0,0,0.1);
    }
    
    .metric-value {
        font-size: 2rem;
        font-weight: 700;
        background: linear-gradient(135deg, #1e3c72 0%, #2a5298 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
    }
    
    .metric-label {
        color: #666;
        font-size: 0.9rem;
        margin-top: 0.5rem;
    }
    
    /* Buttons */
    .stButton > button {
        background: linear-gradient(135deg, #1e3c72 0%, #2a5298 100%);
        color: white;
        border: none;
        padding: 0.75rem 2rem;
        font-weight: 600;
        border-radius: 8px;
        transition: all 0.3s;
        width: 100%;
    }
    
    .stButton > button:hover {
        transform: translateY(-2px);
        box-shadow: 0 5px 20px rgba(30, 60, 114, 0.4);
    }
    
    /* File uploader */
    .stFileUploader {
        background: white;
        padding: 2rem;
        border-radius: 12px;
        border: 2px dashed #e0e0e0;
        transition: all 0.3s;
    }
    
    .stFileUploader:hover {
        border-color: #2a5298;
        background: #fafbff;
    }
    
    /* Success/Warning messages */
    .success-msg {
        background: linear-gradient(135deg, #84fab0 0%, #8fd3f4 100%);
        color: #1a5f3f;
        padding: 1rem;
        border-radius: 8px;
        margin: 1rem 0;
        font-weight: 500;
    }
    
    .warning-msg {
        background: linear-gradient(135deg, #ffeaa7 0%, #fdcb6e 100%);
        color: #8b5a00;
        padding: 1rem;
        border-radius: 8px;
        margin: 1rem 0;
        font-weight: 500;
    }
    
    .info-msg {
        background: linear-gradient(135deg, #a8d8ff 0%, #86c5ff 100%);
        color: #004085;
        padding: 1rem;
        border-radius: 8px;
        margin: 1rem 0;
        font-weight: 500;
    }
    
    /* Select boxes */
    .stSelectbox > div > div {
        background: white;
        border-radius: 8px;
        border: 1px solid #e0e0e0;
    }
    
    /* Data table */
    .dataframe {
        border: none !important;
        border-radius: 8px;
        overflow: hidden;
    }
    
    .dataframe thead {
        background: linear-gradient(135deg, #1e3c72 0%, #2a5298 100%);
        color: white;
    }
    
    .dataframe tbody tr:hover {
        background: #f8f9fa;
    }
    
    /* Receipt table */
    .receipt-table {
        width: 100%;
        border-collapse: collapse;
        margin-top: 1rem;
    }
    
    .receipt-table th {
        background: linear-gradient(135deg, #1e3c72 0%, #2a5298 100%);
        color: white;
        padding: 0.75rem;
        text-align: left;
        font-weight: 600;
    }
    
    .receipt-table td {
        padding: 0.75rem;
        border-bottom: 1px solid #e0e0e0;
    }
    
    .receipt-table tr:hover {
        background: #f8f9fa;
    }
    
    /* Action buttons in table */
    .action-button {
        background: #dc3545;
        color: white;
        border: none;
        padding: 0.25rem 0.5rem;
        border-radius: 4px;
        font-size: 0.8rem;
        cursor: pointer;
    }
    
    .action-button:hover {
        background: #c82333;
    }
</style>
""", unsafe_allow_html=True)

# ============================================
# INITIALIZE SESSION STATE
# ============================================
if 'receipts_db' not in st.session_state:
    st.session_state.receipts_db = []

if 'invoices_db' not in st.session_state:
    st.session_state.invoices_db = []

if 'payment_instructions_db' not in st.session_state:
    st.session_state.payment_instructions_db = []

# ============================================
# HEADER WITH LOGO
# ============================================
st.markdown("""
<div class="app-header">
    <div class="logo-container">
        <img src="https://raw.githubusercontent.com/Daniel43450/aiolos-excel-app/main/Capture.PNG" alt="Aiolos Logo">
    </div>
    <div class="header-text">
        <h1 class="app-title">Aiolos Financial Tools</h1>
        <p class="app-subtitle">Smart financial management made simple</p>
    </div>
</div>
""", unsafe_allow_html=True)

# ============================================
# HELPER FUNCTIONS
# ============================================

# --- DOCX → PDF helper (tries docx2pdf, then LibreOffice) ---
def docx_bytes_to_pdf_bytes(docx_bytes: bytes):
    import tempfile, os, shutil, subprocess
    tmpdir = tempfile.mkdtemp()
    try:
        docx_path = os.path.join(tmpdir, "tmp.docx")
        with open(docx_path, "wb") as f:
            f.write(docx_bytes)

        # ניסיון 1: docx2pdf (Windows/Mac, וגם חלק מסביבות לינוקס)
        try:
            import docx2pdf
            pdf_path = os.path.join(tmpdir, "tmp.pdf")
            docx2pdf.convert(docx_path, pdf_path)
            with open(pdf_path, "rb") as f:
                return f.read()
        except Exception:
            pass

        # ניסיון 2: LibreOffice headless
        try:
            cmd = ["soffice", "--headless", "--convert-to", "pdf", docx_path, "--outdir", tmpdir]
            subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
            pdf_path = os.path.join(tmpdir, "tmp.pdf")
            with open(pdf_path, "rb") as f:
                return f.read()
        except Exception:
            return None
    finally:
        shutil.rmtree(tmpdir, ignore_errors=True)


def find_all_plots(description):
    """Find all plot references in description"""
    PLOTS = [
        'Y1', 'Y2', 'Y3', 'Y6', 'Y4-7', 'Y8', 'R2', 'R4', 'B5', 'G2',
        'R5A', 'R5B', 'R5C', 'R5D', 'W2', 'W8', 'B6', 'G1', 'G12', 'G13', 'B9-10-11'
    ]
    found = []
    for plot in PLOTS:
        if re.search(rf"(?<!\\w){re.escape(plot)}(?!\\w)", description):
            found.append(plot)
    return found

# ============================================
# VILLA OWNERS DATABASE
# ============================================
VILLA_OWNERS = {
    ("Y1", "Villa 1"): "George Bezerianos",
    ("Y1", "Villa 2"): "Shelley Furman Assa",
    ("Y2", "Villa 1"): "Shimrit Bourla",
    ("Y2", "Villa 2"): "Chen Arad",
    ("Y3", "Villa 1"): "Ronen Doron Aviram",
    ("Y3", "Villa 2"): "Ronen Ofec",
    ("Y3", "Villa 3"): "Ran Hai",
    ("Y3", "Villa 4"): "Pepel Weitz Yaniva",
    ("Y3", "Villa 5"): "Eliyahu Ovadia",
    ("Y4-7", "Villa 9"): "Elad Shimon Nissenholtz",
    ("Y4-7", "Villa 10"): "Dan Dikanoff",
    ("G2", "Villa 1"): "Ester Danziger",
    ("G2", "Villa 2"): "Gil Bar el",
    ("G2", "Villa 3"): "Michael & Tal Gurevich",
    ("G2", "Villa 4"): "Alexander Gurevich",
    ("G2", "Villa 5"): "Linkova Oksana M",
    ("G2", "Villa 6"): "Ofir Laor",
    ("G2", "Villa 7"): "",
    ("G2", "Villa 8"): "Patrice Daniel Giami",
    ("G13", "Villa 1"): "Nir Goldberg",
    ("G13", "Villa 2"): "Nir Goldberg",
    ("G13", "Villa 3"): "Keren Goldberg",
    ("G13", "Villa 4"): "Keren Goldberg",
    ("G13", "Villa 5"): "Rachel Goldberg Keidar",
    ("R4", "Villa 1"): "Itach Eyal & Ella",
    ("R4", "Villa 2"): "Nirit & Ofer Mizrahi",
}

# ============================================
# DIAKOFTI PROCESSING FUNCTION
# ============================================
def process_diakofti_file(df):
    """Process Diakofti format files"""
    df = df.dropna(subset=['ΠΕΡΙΓΡΑΦΗ'])
    df['ΠΟΣΟ'] = df['ΠΟΣΟ'].astype(str).str.replace('.', '').str.replace(',', '.').astype(float)
    
    results = []
    for _, row in df.iterrows():
        original_desc = str(row['ΠΕΡΙΓΡΑΦΗ'])
        desc = original_desc.upper()
        amount = abs(row['ΠΟΣΟ'])
        plots = find_all_plots(desc)
        
        if len(plots) == 1:
            plot_val = plots[0]
        elif len(plots) > 1:
            plot_val = "Multiple"
        else:
            plot_val = "All Plots"
        
        is_income = row['ΠΟΣΟ'] > 0
        
        entry = {
            "Date": row['ΗΜ/ΝΙΑ ΚΙΝΗΣΗΣ'],
            "Income/outcome": "Income" if is_income else "Outcome",
            "Plot": plot_val,
            "Expenses Type": "Soft Cost",
            "Type": "",
            "Supplier": "",
            "Description": desc,
            "In": amount if is_income else "",
            "Out": -amount if not is_income else "",
            "Total": amount if is_income else -amount,
            "Progressive Ledger Balance": "",
            "Payment details": "",
            "Original Description": original_desc
        }
        
        filled = False
        
        # ============================================
        # 🔴 Diakofti RULES - ADD YOUR RULES HERE
        
        if "COM POI" in desc or "COM POO" in desc:
            entry["Type"] = "Bank"
            entry["Supplier"] = "Bank"
            entry["Description"] = "Bank fees"
            filled = True

        if "ELECTRICAL INSTALLATION" in desc.upper():
            entry["Type"] = "Electricity"
            entry["Supplier"] = "Engineer"
            entry["Description"] = "Construction works"
            filled = True

        if amount == 1006.77:
            entry["Type"] = "Operation cost"
            entry["Supplier"] = "Worker 1"
            entry["Description"] = "Salary"
            filled = True

        if "ΠΛΗΡΩΜΗ ΒΕΒΑΙΩΜΕΝΕΣ ΣΤΙΣ Δ.Ο.Υ. ΟΦΕΙΛΕΣ" in desc.upper() and amount == 100.16:
            entry["Type"] = "Operation cost"
            entry["Supplier"] = "Worker 1"
            entry["Description"] = "TEKA"
            filled = True

        if "STAVROU" in desc.upper() or "SKANDIA" in desc.upper() or "PLATANOS" in desc.upper() or "TO LIMAN KYTHI GR" in desc.upper():
            entry["Type"] = "General"
            entry["Supplier"] = "F&B"
            entry["Description"] = "F&B"
            filled = True

        if "RF549086180000334044" in desc.upper():
            entry["Plot"] = "G2"
            entry["Expenses Ty"] = "Soft Cost"
            entry["Type"] = "Utility Bills"
            entry["Supplier"] = "Municipality"
            entry["Description"] = "Electricity"
            filled = True

        if "MOREAS S" in desc.upper():
            entry["Type"] = "Project management"
            entry["Supplier"] = "Transportation"
            entry["Description"] = "Transportation"
            filled = True
   
        if "AIOLOS DIAKOFTI EKMETALLEFSI AKINIT" in desc.upper():
            entry["Location"] = "Diakofti"
            entry["Type"] = "Aiolos Diakofti"
            entry["Supplier"] = "Operation cost"
            entry["Description"] = "Reimbursement of expenses"
            filled = True
            
        if "CANVA" in desc.upper():
            entry["Type"] = "General"
            entry["Supplier"] = "Office expenses"
            entry["Description"] = "Office expense"
            entry["Payment details"] = "CANVA"  
            filled = True

        if "ARID" in desc.upper():
            entry["Type"] = "Architect"
            entry["Supplier"] = "ARID"
            entry["Description"] = "Planning"
            filled = True

        if "CLAUDE" in desc.upper():
            entry["Type"] = "General"
            entry["Supplier"] = "Claude AI"
            entry["Description"] = "Office expense"
            filled = True


        if "HERTZ" in desc.upper():
            entry["Type"] = "Project management"
            entry["Supplier"] = "Panayotis"
            entry["Description"] = "Car rent fees"
            filled = True

        if "AVIS" in desc.upper():
            entry["Type"] = "Project management"
            entry["Supplier"] = "Panayotis"
            entry["Description"] = "Car rent fees"
            filled = True

        if "SIXT" in desc.upper():
            entry["Type"] = "Project management"
            entry["Supplier"] = "Panayotis"
            entry["Description"] = "Car rent fees"
            filled = True

        if "PANAYOTIS" in desc.upper():
            entry["Type"] = "Project management"
            entry["Supplier"] = "Panayotis"
            entry["Description"] = "Car rent fees"
            filled = True


        if "EDEN" in desc:
            entry["Type"] = "Project management"
            entry["Supplier"] = "Accommodation"
            entry["Description"] = "Hotel"
            filled = True

        if any(keyword in desc.upper() for keyword in ["TRANSPORT KALLI GR","GRIGORAK KYTHI GR", "O MAGOS KYTHI GR", "STAMATIS KYTHI GR", "KONTOLEO KYTHI GR", "VITSIO KYTHI GR", "BOURNAKI KYTHI GR", "STAVROU KYTHI GR"]):
            entry["Type"] = "General"
            entry["Supplier"] = "F&B"
            entry["Description"] = "F&B"
            filled = True


        if "ALL PLOTS MARKETING" in desc:
            entry["Type"] = "Marketing"
            entry["Supplier"] = "Marketing"
            entry["Description"] = "Marketing Services fee"
            filled = True

        if "CALEN" in desc or "HARD COST" in desc:
            entry["Expenses Type"] = "Hard Cost"
            entry["Type"] = "Contractor"
            entry["Supplier"] = "Calen"
            entry["Description"] = "Construction works"
            filled = True

        if "SUPERVISION" in desc:
            entry["Type"] = "Supervision"
            entry["Supplier"] = "TAG ARCHITECTS"
            entry["Description"] = "Supervision"
            filled = True

        if "HOLIDAYS TEL" in desc:
            entry["Type"] = "Project management"
            entry["Supplier"] = "Transportation"
            entry["Description"] = "Flight"
            filled = True

        if "EL AL" in desc:
            entry["Type"] = "Project management"
            entry["Supplier"] = "Transportation"
            entry["Description"] = "Flight"
            filled = True

        if any(keyword in desc.upper() for keyword in ["FACEBOOK", "FACEBK", "FB.ME", "META"]):
            entry["Type"] = "Marketing"
            entry["Supplier"] = "Marketing"
            entry["Description"] = "Marketing Services fee"
            filled = True


        if any(term in desc for term in ["ACCOUNTING", "BOOKKEEP", "ECOVIS"]) and not any(word in desc for word in ["YAG", "TAG"]):
            entry["Type"] = "Accounting"
            entry["Supplier"] = "Ecovis"
            entry["Description"] = "Accountant monthly fees"
            filled = True

        if "GAS" in desc:
            entry["Type"] = "Project management"
            entry["Supplier"] = "Transportation"
            entry["Description"] = "Gas station"
            filled = True

        if "DRAKAKIS" in desc:
            entry["Type"] = "Project management"
            entry["Supplier"] = "Drakakis Tours"
            entry["Description"] = "Car rent fees"
            filled = True

        if "FLIGHT" in desc or "AEGEAN" in desc:
            entry["Type"] = "Project management"
            entry["Supplier"] = "Transportation"
            entry["Description"] = "Flight"
            filled = True

        if "TONY S" in desc or "Tony S" in desc or "tony s" in desc or "eat" in desc or "EAT" in desc:
            entry["Type"] = "General"
            entry["Supplier"] = "F&B"
            entry["Description"] = "F&B"
            filled = True


        if any(word in desc for word in ["AEGEANWEB", "AEGEAN", "OLYMPIC", "SKY", "ISRAIR", "WIZZ"]):
            entry["Type"] = "Project management"
            entry["Supplier"] = "Transportation"
            entry["Description"] = "Flight"
            filled = True

        if any(word in desc for word in ["DINNER", "FOOD", "CAFE", "COFFEE", "LUNCH", "BREAKFAST"]):
            entry["Type"] = "General"
            entry["Supplier"] = "F&B"
            entry["Description"] = "F&B"
            filled = True

        if ("broker" in desc or "Broker" in desc or "BROKER" in desc) and (
            "villa 1" in desc or "Villa 1" in desc or "VILLA 1" in desc):
            entry["Type"] = "Brokers"
            entry["Supplier"] = "Buyer Villa 1"
            entry["Description"] = "Broker fees"
            filled = True

        if ("broker" in desc or "Broker" in desc or "BROKER" in desc) and (
            "villa 2" in desc or "Villa 2" in desc or "VILLA 2" in desc):
            entry["Type"] = "Brokers"
            entry["Supplier"] = "Buyer Villa 2"
            entry["Description"] = "Broker fees"
            filled = True

        if "RF919086180000334" in desc:
            entry["Plot"] = "R4"
            entry["Expenses Type"] = "Soft Cost"
            entry["Type"] = "Utility Bills"
            entry["Supplier"] = "Municipality"
            entry["Description"] = "Electricity"
            filled = True


        if ("broker" in desc or "Broker" in desc or "BROKER" in desc) and (
            "villa 3" in desc or "Villa 3" in desc or "VILLA 3" in desc):
            entry["Type"] = "Brokers"
            entry["Supplier"] = "Buyer Villa 3"
            entry["Description"] = "Broker fees"
            filled = True

        if ("broker" in desc or "Broker" in desc or "BROKER" in desc) and (
            "villa 4" in desc or "Villa 4" in desc or "VILLA 4" in desc):
            entry["Type"] = "Brokers"
            entry["Supplier"] = "Buyer Villa 4"
            entry["Description"] = "Broker fees"
            filled = True

        if ("broker" in desc or "Broker" in desc or "BROKER" in desc) and (
            "villa 5" in desc or "Villa 5" in desc or "VILLA 5" in desc):
            entry["Type"] = "Brokers"
            entry["Supplier"] = "Buyer Villa 5"
            entry["Description"] = "Broker fees"
            filled = True

        if ("broker" in desc or "Broker" in desc or "BROKER" in desc) and (
            "villa 6" in desc or "Villa 6" in desc or "VILLA 6" in desc):
            entry["Type"] = "Brokers"
            entry["Supplier"] = "Buyer Villa 6"
            entry["Description"] = "Broker fees"
            filled = True


        if "GOOGLE" in desc:
            entry["Type"] = "Marketing"
            entry["Supplier"] = "Marketing"
            entry["Description"] = "Marketing Services fee"
            filled = True

        if "CRM" in desc:
            entry["Type"] = "Marketing"
            entry["Supplier"] = "reWire"
            entry["Description"] = "CRM"
            filled = True

        if "RF91908618000033404472101" in desc or "PROT-RF549086180000334" in desc:
            entry["Type"] = "Utility Bills"
            entry["Supplier"] = "Municipality"
            entry["Description"] = "Electricity"
            entry["Plot"] = "G2"
            filled = True

        if "RF38908618000033404445701" in desc or "RF389086180000334044" in desc:  
            entry["Type"] = "Utility Bills"
            entry["Supplier"] = "Municipality"
            entry["Description"] = "Electricity"
            entry["Plot"] = "Y3"
            filled = True

        if "RF91908618000033404472101" in desc or "PROT-919086180000334" in desc:
            entry["Type"] = "Utility Bills"
            entry["Supplier"] = "Municipality"
            entry["Description"] = "Electricity"
            entry["Plot"] = "R4"
            filled = True

        if "UBER" in desc or "TAXI" in desc:
            entry["Type"] = "Project management"
            entry["Supplier"] = "Transportation"
            entry["Description"] = "Athens Taxi"
            filled = True

        if "BEAUTIFU SAN" in desc.upper():
            entry["Type"] = "General"
            entry["Supplier"] = "BEAUTIFUL"
            entry["Description"] = "Office expense"
            filled = True

        if "OPENAI" in desc:
            entry["Type"] = "General"
            entry["Supplier"] = "Office expenses"
            entry["Description"] = "Office expense"
            filled = True

        if "TAG" in desc:
            entry["Type"] = "Architect"
            entry["Supplier"] = "TAG ARCHITECTS"
            if "SUP" in desc:
                entry["Description"] = "Supervision"
            else:
                entry["Description"] = "Planning"
            filled = True

        if "OASA" in desc:
            entry["Type"] = "Project management"
            entry["Supplier"] = "Transportation"
            entry["Description"] = "Transportation"
            filled = True

        if "ΔΗΜΟ-RF369029090000097" in desc:
            entry["Type"] = "Utility Bills"
            entry["Supplier"] = "Municipality"
            entry["Description"] = "Water"
            entry["Plot"] = "Y3"
            filled = True


        if any(term in desc for term in ["MANAGEMENT", "MANAG.", "MGMT", "MNGMT"]) and row['ΠΟΣΟ'] in [-1550, -1550.00, -1550.0, 1550.00, 1550.0, 2055, 2055.0, 2057.0, 1550]:
            entry["Type"] = "Worker 1"
            entry["Supplier"] = "Aiolos Athens"
            entry["Description"] = "management fees"
            filled = True

        if any(term in desc for term in ["COSM", "COSMOTE", "PHONE"]):
            entry["Type"] = "Utility Bills"
            entry["Supplier"] = "Cosmote"
            entry["Description"] = "Phone bill"
            filled = True

        if "RF389086180000334" in desc:
            entry["Type"] = "Utility Bills"
            entry["Supplier"] = "Municipality"
            entry["Description"] = "Electricity"
            entry["Plot"] = "Y3"
            filled = True



        # END OF Diakofti RULES
        # ============================================
        
        if not filled:
            entry["Description"] = f"🟨 {entry['Description']}"
        
        results.append(entry)
    
    return pd.DataFrame(results)

# ============================================
# Athens PROCESSING FUNCTION
# ============================================
def process_athens_file(df):
    """Process Athens format files"""
    df = df.copy()
    df['Ημερομηνία'] = pd.to_datetime(df['Ημερομηνία'], dayfirst=True, errors='coerce')
    df = df.dropna(subset=['Περιγραφή'])
    
    results = []
    for _, row in df.iterrows():
        original_desc = str(row['Περιγραφή'])
        desc = original_desc.upper()
        amount = abs(float(str(row['Ποσό συναλλαγής']).replace(',', '.')))
        
        entry = {
            "Date": row['Ημερομηνία'].strftime('%d/%m/%Y') if not pd.isnull(row['Ημερομηνία']) else '',
            "Income/Outcome": "Income" if row['Ποσό εντολής'] > 0 else "Outcome",
            "Expenses Type": "Soft Cost",
            "Location": "All Projects",
            "Project": "All Projects",
            "Supplier": "",
            "Type": "",
            "Description": desc,
            "Income": amount if row['Ποσό εντολής'] > 0 else "",
            "Outcome": -amount if row['Ποσό εντολής'] < 0 else "",
            "Total": amount if row['Ποσό εντολής'] > 0 else -amount,
            "Balance": "",
            "Repayment": "",
            "Original Description": original_desc
        }
        
        filled = False
        
        # ============================================
        # 🔵 Athens RULES - ADD YOUR RULES HERE
        if any(word in desc for word in ["DINNER", "FOOD", "CAFE", "COFFEE", "LUNCH", "BREAKFAST", "ΦΑΓΗΤΟ", "ΕΣΤΙΑΤΟΡΙΟ", "ΚΑΦΕ"]):
            entry["Type"] = "F&B"
            entry["Supplier"] = "General"
            entry["Description"] = "F&B"
            filled = True

        if "TEKA" in desc and round(amount, 2) == 76.66:
            entry["Supplier"] = "Worker 1"
            entry["Type"] = "Operation cost"
            entry["Description"] = "TEKA"
            filled = True

        if "EPASSNAODOSGR" in desc.upper():
            entry["Supplier"] = "Transportation"
            entry["Type"] = "General"
            entry["Description"] = "Toll road"
            filled = True

        if any(keyword in desc.upper() for keyword in ["FACEBOOK", "FACEBK", "FB.ME", "META"]):
            entry["Supplier"] = "Marketing"
            entry["Type"] = "Marketing"
            entry["Description"] = "Marketing Services fee"
            filled = True

        if "BAGELDB" in desc:
            entry["Type"] = "Marketing"
            entry["Supplier"] = "BagelDB"
            entry["Description"] = "Website"
            filled = True

        if any(variant in desc for variant in ["AP MICHALOPOULOS SIA", "Ap Michalopoulos Sia", "ap michalopoulos sia"]):
            entry["Type"] = "F&B"
            entry["Supplier"] = "General"
            entry["Description"] = "F&B"
            filled = True
           
        if any(word in desc for word in ["AVIS", "HERTZ", "SIXT", "CAR RENTAL"]):
            entry["Type"] = "Transportation"
            entry["Supplier"] = "Transportation"
            entry["Description"] = "Car rental"
            entry["Plot"] = "All Projects"
            filled = True
           
        if "COSMOTE" in desc:
            entry["Location"] = "Mobee"
            entry["Project"] = "Mobee"
            entry["Supplier"] = "Cosmote"
            entry["Type"] = "Project Management"
            entry["Description"] = "Office expenses"
            filled = True

           
        if any(word in desc for word in ["BAKERY","KENTRIKI ENOSI EPIME", "CAFFE", "CAFE", "EAT", "BEVERAGE", "PIZA", "BURGER"]):
            entry["Type"] = "F&B"
            entry["Supplier"] = "General"
            entry["Description"] = "F&B"
            filled = True

        if any(keyword in desc for keyword in ["WEBCCDOMAINCOM", "webccdomaincom", "Webccdomaincom"]):
            entry["Type"] = "Marketing"
            entry["Supplier"] = "BagelDB"
            entry["Description"] = "Website"
            entry["Repayment"] = "DOMAIN"
            filled = True
           
        if round(amount, 2) == 496.00 and any(word in desc for word in ["ECOVIS", "FEE", "FEES"]):
            entry["Supplier"] = "Accountant"
            entry["Type"] = "Ecovis"
            entry["Description"] = "Accountant monthly fees"
            filled = True

        if abs(row['Ποσό συναλλαγής']) == 256.41 and row['Ποσό συναλλαγής'] < 0:
            entry["Type"] = "Tax"
            entry["Supplier"] = "Authorities"
            entry["Description"] = "EFKA"
            entry["Repayment"] = "UDI EFKA"
            filled = True
           
        if "MANAGEMENT FEE" in desc or row['Ποσό συναλλαγής'] in [-1810, 1810, -1810.00, 1810.00]:
            entry["Type"] = "Mobee Management"
            entry["Supplier"] = "Konstantinos"
            entry["Description"] = "Management fee"
            filled = True

        if "ΠΡΟΜΗΘΕΙΕΣ ΕΞΟΔΑ" in desc and amount <= 5:
            entry["Type"] = "Bank fees"
            entry["Supplier"] = "Bank"
            entry["Description"] = "Bank fees"
            filled = True

        if "AIOLOS DIAKOFTI" in desc and 1520 <= amount <= 1570:
            entry["Supplier"] = "Aiolos Diakofti"
            entry["Type"] = "Operation cost"
            entry["Description"] = "Reimbursement of expenses"
            filled = True

        if ("ΚΑΛΛΙΦΡΟΝΑ 3" in desc or "ΚΑΛΛΙΦΡΟΝΑ3" in desc) and row['Ποσό συναλλαγής'] > 0:
            entry["Type"] = "Mobee Management"
            entry["Supplier"] = "Kalliforna"
            entry["Description"] = "Management fee"
            entry["Location"] = "Mobee"
            filled = True
           
        if "ΠΛΗΡΩΜΗ ΕΦΚΑ ΕΡΓΟΔΟΤΙΚΕΣ ΕΙΣΦΟΡΕΣ" in desc:
            entry["Type"] = "Tax"
            entry["Supplier"] = "Authorities"
            entry["Description"] = "EFKA"
            filled = True

        if "PLAKENTIA" in desc:
            entry["Type"] = "Transportation"
            entry["Supplier"] = "General"
            entry["Description"] = "Metro"
            filled = True

        if "MICROSOFT" in desc:
            entry["Type"] = "Project Management"
            entry["Supplier"] = "Microsoft"
            entry["Description"] = "Office expenses"
            filled = True

        if "LEFKES VILLAS PROJECT MONOPROSOPI" in desc:
            entry["Supplier"] = "Lefkes Villas"
            entry["Type"] = "Project Management"
            entry["Description"] = "Management fee"

        if "LEFKES" in desc:
            entry["Location"] = "Lefkes"

        if "BEN SHAHAR" in desc:
            entry["Supplier"] = "Ben Shahar"
            entry["Type"] = "Project Management"
            entry["Description"] = "Management fee"
            entry["Location"] = "Lefkes"
                 
        if "PARKING" in desc:
            entry["Type"] = "Transportation"
            entry["Supplier"] = "Parking"
            entry["Description"] = "Parking"
            filled = True

        if "kallifrona 3 ekmetalleysi akiniton" in desc.lower() and amount == 2450:
            entry["Location"] = "Mobee"
            entry["Project"] = "Mobee"
            entry["Supplier"] = "Kalliforna"
            entry["Type"] = "Mobee Management"
            entry["Description"] = "Management fee"
            filled = True

        if "ECOVIS" in desc:
            entry["Type"] = "Ecovis"
            entry["Supplier"] = "Accountant"
            entry["Description"] = "Accountant monthly fees"
            filled = True

        if row['Ποσό συναλλαγής'] == 4960:
            entry["Type"] = "Project Management"
            entry["Supplier"] = "Lefkes Villas"
            entry["Description"] = "Management fee"
            entry["Location"] = "Lefkes"
            entry["Expenses Type"] = "Soft cost"
            filled = True

        if "HAREL" in desc:
            entry["Type"] = "Project Management"
            entry["Supplier"] = "General"
            entry["Description"] = "Office expenses"
            filled = True

        if "SHELL" in desc:
            entry["Type"] = "Transportation"
            entry["Supplier"] = "General"
            entry["Description"] = "Gas station"
            filled = True

        if "OASA" in desc:
            entry["Type"] = "Transportation"
            entry["Supplier"] = "General"
            entry["Description"] = "Metro"
            filled = True

        if "WORKER 1" in desc:
            entry["Type"] = "Operation cost"
            entry["Supplier"] = "Worker 1"
            entry["Description"] = "Salary"
            filled = True

        if any(word in desc for word in ["AEGEANWEB", "AEGEAN", "OLYMPIC", "SKY", "ISRAIR", "WIZZ"]):
            entry["Type"] = "Transportation"
            entry["Supplier"] = "General"
            entry["Description"] = "Flight"
            filled = True

        if ("ΠΛΗΡΩΜΗ ΒΕΒΑΙΩΜΕΝΕΣ ΣΤΙΣ Δ.Ο.Υ. ΟΦΕΙΛΕΣ" in desc and (amount == 76.66 or amount == -76.66)):
            entry["Supplier"] = "Greek Tax Office"
            entry["Type"] = "Tax Payment"
            entry["Description"] = "DOY Confirmed Payment"
            filled = True

        if "PARKAROUND" in desc:
            entry["Type"] = "Transportation"
            entry["Supplier"] = "Parking"
            entry["Description"] = "Parking"
            filled = True

        if any(word in desc for word in ["ATTIKI"]):
            entry["Type"] = "Transportation"
            entry["Supplier"] = "General"
            entry["Description"] = "Toll road"
            filled = True

        if any(word in desc for word in ["UBER", "UBR"]):
            entry["Type"] = "Transportation"
            entry["Supplier"] = "General"
            entry["Description"] = "Uber"
            filled = True

        if "GOOGLE" in desc:
            entry["Type"] = "Marketing"
            entry["Supplier"] = "Google"
            entry["Description"] = "Campaign"
            filled = True

        if "PETRELION" in desc:
            entry["Type"] = "Transportation"
            entry["Supplier"] = "General"
            entry["Description"] = "Gas station"
            filled = True

        if any(word in desc for word in ["ΠΡΟΜΗΘ", "ΜΗΝ", "ΠΑΡ", "ΕΞΟΔΑ"]) and amount <= 5:
            entry["Type"] = "Bank"
            entry["Supplier"] = "Bank"
            entry["Description"] = "Bank fees"
            filled = True


        # END OF Athens RULES
        # ============================================
        
        if not filled:
            entry["Description"] = f"🟨 {entry['Description']}"
        
        results.append(entry)
    
    result_df = pd.DataFrame(results)
    
    # Reorder columns
    column_order = [
        "Date", "Income/Outcome", "Expenses Type", "Location", "Project",
        "Supplier", "Type", "Description", "Income", "Outcome", "Total",
        "Balance", "Repayment", "Original Description"
    ]
    
    return result_df[column_order]

# ============================================
# Ilisia NBG PROCESSING FUNCTION
# ============================================
def process_ilisia_file(df):
    """Process Ilisia NBG format files (robust to column names)"""
    df = df.copy()

    # --- resolve columns (date/desc/amount can arrive with several names) ---
    def pick(*names):
        for n in names:
            if n in df.columns:
                return n
        return None

    col_date   = pick('ΗΜ/ΝΙΑ ΚΙΝΗΣΗΣ', 'Ημερομηνία', 'Valeur')
    col_desc   = pick('ΠΕΡΙΓΡΑΦΗ', 'Περιγραφή')
    col_amount = pick('ΠΟΣΟ', 'Ποσό εντολής', 'Ποσό συναλλαγής')


    if col_desc is None or col_amount is None:
        raise ValueError("Ilisia: description or amount column missing")

    # --- normalize canonical columns expected by your RULES ---
    # Description -> ΠΕΡΙΓΡΑΦΗ
    if col_desc != 'ΠΕΡΙΓΡΑΦΗ':
        df['ΠΕΡΙΓΡΑΦΗ'] = df[col_desc]

    # Date -> ΗΜ/ΝΙΑ ΚΙΝΗΣΗΣ  (parse then keep canonical)
    if col_date:
        df[col_date] = pd.to_datetime(df[col_date], dayfirst=True, errors='coerce')
        if col_date != 'ΗΜ/ΝΙΑ ΚΙΝΗΣΗΣ':
            df['ΗΜ/ΝΙΑ ΚΙΝΗΣΗΣ'] = df[col_date]
    else:
        df['ΗΜ/ΝΙΑ ΚΙΝΗΣΗΣ'] = pd.NaT

    # Amount -> ΠΟΣΟ  (keep sign, handle commas/dots)
    if pd.api.types.is_numeric_dtype(df[col_amount]):
        df['_signed_amount'] = df[col_amount].astype(float)
    else:
        df['_signed_amount'] = (
            df[col_amount].astype(str)
            .str.replace('.', '', regex=False)     # thousands dot
            .str.replace(',', '.', regex=False)    # decimal comma
            .str.replace(r'[^\d\.\-]', '', regex=True)
            .replace({'': '0', '-': '0'})
            .astype(float)
        )
    if col_amount != 'ΠΟΣΟ':
        df['ΠΟΣΟ'] = df['_signed_amount']
    else:
        # ensure ΠΟΣΟ is clean float even if read as string
        if not pd.api.types.is_numeric_dtype(df['ΠΟΣΟ']):
            df['ΠΟΣΟ'] = (
                df['ΠΟΣΟ'].astype(str)
                .str.replace('.', '', regex=False)
                .str.replace(',', '.', regex=False)
                .str.replace(r'[^\d\.\-]', '', regex=True)
                .replace({'': '0', '-': '0'})
                .astype(float)
            )

    # clean rows
    df = df.dropna(subset=['ΠΕΡΙΓΡΑΦΗ'])

    results = []
    for _, row in df.iterrows():
        original_desc = str(row['ΠΕΡΙΓΡΑΦΗ'])
        desc = original_desc.upper()
        amt = float(row['ΠΟΣΟ'])
        amount = abs(amt)

        # plot detect
        plots = find_all_plots(desc)
        if len(plots) == 1:
            plot_val = plots[0]
        elif len(plots) > 1:
            plot_val = "Multiple"
        else:
            plot_val = "G1 - Manolis"

        is_income = amt > 0

        entry = {
            "Date": row['ΗΜ/ΝΙΑ ΚΙΝΗΣΗΣ'].strftime('%d/%m/%Y') if pd.notnull(row['ΗΜ/ΝΙΑ ΚΙΝΗΣΗΣ']) else '',
            "Income/outcome": "Income" if is_income else "Outcome",
            "Plot": plot_val,
            "Expenses Type": "Soft Cost",
            "Type": "",
            "Supplier": "",
            "Description": desc,
            "In": amount if is_income else "",
            "Out": -amount if not is_income else "",
            "Vat": "",
            "Total": amount if is_income else -amount,
            "Progressive Ledger Balance": "",
            "Payment details": "",
            "Original Description": original_desc,
            "Year": row['ΗΜ/ΝΙΑ ΚΙΝΗΣΗΣ'].year if pd.notnull(row['ΗΜ/ΝΙΑ ΚΙΝΗΣΗΣ']) else "",
            "Bank": "NBG" 
            
        }

        filled = False
        
        # ============================================
        # 🔴 Ilisia RULES - ADD YOUR RULES HERE
        # ============================================
        
        if "COM POI" in desc or "COM POO" in desc:
            entry["Type"] = "Bank"
            entry["Supplier"] = "Bank"
            entry["Description"] = "Bank fees"
            filled = True

        
        # --- Rule: Booking Operation Income (positive) ---
        if ("ΠΚ/00505341795" in desc or "ΠΚ/02505341795" in desc) and row['ΠΟΣΟ'] > 0:
            entry["Expenses Type"] = "Operation Income"
            entry["Type"] = "Accommodation"
            entry["Supplier"] = "Booking"
            entry["Description"] = "Accommodation fees"
            filled = True

        # --- Rule: Booking Refund (negative) ---
        if ("ΠΚ/00505341795" in desc or "ΠΚ/02505341795" in desc) and row['ΠΟΣΟ'] < 0:
            entry["Expenses Type"] = "Operation Income"
            entry["Type"] = "Accommodation"
            entry["Supplier"] = "Booking"
            entry["Description"] = "Booking refund"
            filled = True

        # --- Rule: Airbnb income ---
        if "AIRBNB" in desc and row['ΠΟΣΟ'] > 0:
            entry["Expenses Type"] = "Operation Income"
            entry["Type"] = "Accommodation"
            entry["Supplier"] = "Booking"
            entry["Description"] = "Accommodation fees"
            filled = True

        # --- Rule: Airbnb refund ---
        if "AIRBNB" in desc and row['ΠΟΣΟ'] < 0:
            entry["Expenses Type"] = "Operation Income"
            entry["Type"] = "Accommodation"
            entry["Supplier"] = "Booking"
            entry["Description"] = "Booking refund"
            filled = True
            
        if "ΠΚ/02555341795" in desc and row['ΠΟΣΟ'] > 0:
            entry["Expenses Type"] = "Operation Income"
            entry["Type"] = "Accommodation"
            entry["Supplier"] = "Booking"
            entry["Description"] = "Accommodation fees"
            filled = True

        # --- Rule: Loan repayments (LOAN / ΕΝΤΟΛΗ/ΕΜΒΑΣΜΑ / MAGONEZOS) ---
        if any(term in desc for term in ["LOAN", "ΕΝΤΟΛΗ/ΕΜΒΑΣΜΑ ΣΕ ΑΛΛΗ ΤΡΑΠΕΖΑ", "MAGONEZOS EMMANOUIL", "MAGONEZOS"]):
            entry["Expenses Type"] = "Loan"
            entry["Type"] = "Hotel operation"
            entry["Supplier"] = "Loan Broker"
            entry["Description"] = "Loan repayment"
            filled = True    

        # --- Rule: ΠΚ/00215341795 Booking transactions ---
        if "ΠΚ/00215341795" in desc and row['ΠΟΣΟ'] > 0:
            entry["Expenses Type"] = "Operation Income"
            entry["Type"] = "Accommodation"
            entry["Supplier"] = "Booking"
            entry["Description"] = "Accommodation fees"
            filled = True

        if "ΠΚ/00215341795" in desc and row['ΠΟΣΟ'] < 0:
            entry["Expenses Type"] = "Operation Income"
            entry["Type"] = "Accommodation"
            entry["Supplier"] = "Booking"
            entry["Description"] = "Booking refund"
            filled = True


        # --- Rule: ΠΚ/00555341795 Booking transactions ---
        if "ΠΚ/00555341795" in desc and row['ΠΟΣΟ'] > 0:
            entry["Expenses Type"] = "Operation Income"
            entry["Type"] = "Accommodation"
            entry["Supplier"] = "Booking"
            entry["Description"] = "Accommodation fees"
            filled = True

        if "ΠΚ/00555341795" in desc and row['ΠΟΣΟ'] < 0:
            entry["Expenses Type"] = "Operation Income"
            entry["Type"] = "Accommodation"
            entry["Supplier"] = "Booking"
            entry["Description"] = "Booking refund"
            filled = True



        # --- Rule: ΠΚ/00525341795 Booking transactions ---
        if "ΠΚ/00525341795" in desc and row['ΠΟΣΟ'] > 0:
            entry["Expenses Type"] = "Operation Income"
            entry["Type"] = "Accommodation"
            entry["Supplier"] = "Booking"
            entry["Description"] = "Accommodation fees"
            filled = True

        if "ΠΚ/00525341795" in desc and row['ΠΟΣΟ'] < 0:
            entry["Expenses Type"] = "Operation Income"
            entry["Type"] = "Accommodation"
            entry["Supplier"] = "Booking"
            entry["Description"] = "Booking refund"
            filled = True


        if "ΠΚ/02555341795" in desc and row['ΠΟΣΟ'] < 0:
            entry["Expenses Type"] = "Operation Income"
            entry["Type"] = "Accommodation"
            entry["Supplier"] = "Booking"
            entry["Description"] = "Booking refund"
            filled = True
            
        if "ΠΡΟΜΗΘΕΙΑ ΕΝΤΟΛΗΣ" in desc:
            entry["Expenses Type"] = "Soft Cost"
            entry["Type"] = "Bank"
            entry["Supplier"] = "Bank"
            entry["Description"] = "Bank fees"
            filled = True

        # --- Rule: Bank expenses / ΠΡΟΜΗΘΕΙΕΣ ΕΞΟΔΑ ---
        if "ΠΡΟΜΗΘΕΙΕΣ ΕΞΟΔΑ" in desc:
            entry["Expenses Type"] = "Soft Cost"
            entry["Type"] = "Bank"
            entry["Supplier"] = "Bank"
            entry["Description"] = "Bank fees"
            filled = True

    
        if any(k in desc for k in ["PROTERGIA", "ENERGETICA", "DEI", "ΔΕΗ"]):
            entry["Expenses Type"] = "Soft Cost"
            entry["Type"] = "Hotel operation"
            entry["Supplier"] = "Electricity"
            entry["Description"] = "Electricity bill"
            filled = True


        # --- Rule: Pool cleaning invoice ---
        if "INV400009529618476" in desc:
            entry["Expenses Type"] = "Soft Cost"
            entry["Type"] = "Hotel operation"
            entry["Supplier"] = "Cleaning"
            entry["Description"] = "Pool"
            filled = True

        if any(term in desc.upper() for term in ["POOL", "POOLS"]):
            entry["Expenses Type"] = "Soft Cost"
            entry["Type"] = "Hotel operation"
            entry["Supplier"] = "Cleaning"
            entry["Description"] = "Pool"
            filled = True

        

        if any(k in desc for k in ["ΠΡΟΜΗΘΕΙΑ ΕΝΤΟΛΗΣ", "ΕΞΟΔΑ ΤΡ ΠΛΗΡΩΜΗΣ"]):
            entry["Expenses Type"] = "Soft Cost"
            entry["Type"] = "Bank"
            entry["Supplier"] = "Bank"
            entry["Description"] = "Bank fees"
            filled = True



        # --- Rule: Booking.com Accommodation Income (positive) ---
        if "BOOKING.COM B.V." in desc and row['ΠΟΣΟ'] > 0:
            entry["Expenses Type"] = "Operation Income"
            entry["Type"] = "Accommodation"
            entry["Supplier"] = "Booking"
            entry["Description"] = "Accommodation fees"
            filled = True

        # --- Rule: Social Media Invoice 56 ---
        if "SOCIAL MEDIA INV 56" in desc:
            entry["Expenses Type"] = "Soft Cost"
            entry["Type"] = "Hotel operation"
            entry["Supplier"] = "Vassilis"
            entry["Description"] = "Promotion"
            filled = True

        # --- Rule: Transfer Between Accounts August (negative) ---
        if "TRANSFER BETWEEN ACCOUNTS AUGUST" in desc and row['ΠΟΣΟ'] < 0:
            entry["Expenses Type"] = "Soft Cost"
            entry["Type"] = "Cash facilitation"
            entry["Supplier"] = "Hotel"
            entry["Description"] = "Cash facilitation"
            filled = True

        # --- Rule: Septic Tank expenses ---
        if "SEPTIC" in desc:
            entry["Expenses Type"] = "Soft Cost"
            entry["Type"] = "Septic Tank"
            entry["Supplier"] = "Septic Tank"
            entry["Description"] = "Septic Tank"
            filled = True

        # --- Rule: Roompay Invoice Registration ---
        if "ROOMPAY INVOICE REGISTRATION" in desc:
            entry["Expenses Type"] = "Soft Cost"
            entry["Type"] = "Hotel operation"
            entry["Supplier"] = "Web Hotelier"
            entry["Description"] = "Website"
            filled = True

        # --- Rule: Bank Fees (ΠΡΟΜΗΘΕΙΕΣ ΕΞΟΔΑ) ---
        if "ΠΡΟΜΗΘΕΙΕΣ ΕΞΟΔΑ" in desc:
            entry["Expenses Type"] = "Soft Cost"
            entry["Type"] = "Bank"
            entry["Supplier"] = "Bank"
            entry["Description"] = "Bank fees"
            filled = True

        # --- Rule: Etheras Properties Management / Anna Kythira Supervision ---
        if "ETHERAS PROPERTIES MANAGEMENT" in desc and any (term in desc for term in ["LOURANTOU INVOICE", "MANAGEMENT", "SUPERVISION"]):
            entry["Expenses Type"] = "Hotel operation"
            entry["Type"] = "Anna Kythira"
            entry["Supplier"] = "Anna Kythira"
            entry["Description"] = "Supervision monthly fee"
            filled = True
                    

        # --- New Rule 1 ---
        if ("ΠΚ/00505341795" in desc or "ΠΚ/02505341795" in desc) and row['ΠΟΣΟ'] > 0:
            entry["Expenses Type"] = "Operation Income"
            entry["Type"] = "Accommodation"
            entry["Supplier"] = "Booking"
            entry["Description"] = "Accommodation fees"
            filled = True


        if "STAMATIS PANAGIOTIS STAVRO" in desc.upper():
            entry["Plot"] = "G1 - Manolis"
            entry["Expenses Type"] = "Operation Income"
            entry["Type"] = "Rent"
            entry["Supplier"] = "Tenant - Taverne"
            entry["Description"] = "Monthly Taverne rent"
            filled = True

        if any(term in desc.upper() for term in ["TRANSFER BETWEEN ACCOUNTS", "NBG TO EURO"]):
            entry["Plot"] = "G1 - Manolis"
            entry["Expenses Type"] = "Operation Income"
            entry["Type"] = "Accommodation"
            entry["Supplier"] = "Booking"
            entry["Description"] = "Accommodation fees"
            filled = True

        if "ZARA" in desc.upper():
            entry["Type"] = "Hotel operation"
            entry["Supplier"] = "Maintenance"
            entry["Description"] = "Maintenance"
            filled = True

        if "WATT-VOLT" in desc.upper():
            entry["Type"] = "Authorities"
            entry["Supplier"] = "Electricity"
            entry["Description"] = "Electricity"
            filled = True

        if "ARID" in desc.upper():
            entry["Type"] = "Architect"
            entry["Supplier"] = "ARID"
            entry["Description"] = "Planning"
            filled = True

        if "PANAYOTIS" in desc.upper():
            entry["Type"] = "Project management"
            entry["Supplier"] = "Panayotis"
            entry["Description"] = "Car rent fees"
            filled = True

        if "EDEN" in desc:
            entry["Type"] = "Project management"
            entry["Supplier"] = "Accommodation"
            entry["Description"] = "Hotel"
            filled = True

        if any(keyword in desc.upper() for keyword in ["TRANSPORT KALLI GR","GRIGORAK KYTHI GR", "O MAGOS KYTHI GR", "STAMATIS KYTHI GR", "KONTOLEO KYTHI GR", "VITSIO KYTHI GR", "BOURNAKI KYTHI GR", "STAVROU KYTHI GR"]):
            entry["Type"] = "General"
            entry["Supplier"] = "F&B"
            entry["Description"] = "F&B"
            filled = True

        if "ALL PLOTS MARKETING" in desc:
            entry["Type"] = "Marketing"
            entry["Supplier"] = "Marketing"
            entry["Description"] = "Marketing Services fee"
            filled = True

        if "CALEN" in desc or "HARD COST" in desc:
            entry["Expenses Type"] = "Hard Cost"
            entry["Type"] = "Contractor"
            entry["Supplier"] = "Calen"
            entry["Description"] = "Construction works"
            filled = True

        if "SUPERVISION" in desc:
            entry["Type"] = "Supervision"
            entry["Supplier"] = "TAG ARCHITECTS"
            entry["Description"] = "Supervision"
            filled = True

        if "HOLIDAYS TEL" in desc:
            entry["Type"] = "Project management"
            entry["Supplier"] = "Transportation"
            entry["Description"] = "Flight"
            filled = True

        if "EL AL" in desc:
            entry["Type"] = "Project management"
            entry["Supplier"] = "Transportation"
            entry["Description"] = "Flight"
            filled = True

        if any(keyword in desc.upper() for keyword in ["FACEBOOK", "FACEBK", "FB.ME", "META"]):
            entry["Type"] = "Marketing"
            entry["Supplier"] = "Marketing"
            entry["Description"] = "Marketing Services fee"
            filled = True

        if any(term in desc for term in ["ACCOUNTING", "BOOKKEEP", "ECOVIS"]) and not any(word in desc for word in ["YAG", "TAG"]):
            entry["Type"] = "Accounting"
            entry["Supplier"] = "Ecovis"
            entry["Description"] = "Accountant monthly fees"
            filled = True

        if "GAS" in desc:
            entry["Type"] = "Project management"
            entry["Supplier"] = "Transportation"
            entry["Description"] = "Gas station"
            filled = True

        if "DRAKAKIS" in desc:
            entry["Type"] = "Project management"
            entry["Supplier"] = "Drakakis Tours"
            entry["Description"] = "Car rent fees"
            filled = True

        if "FLIGHT" in desc or "AEGEAN" in desc:
            entry["Type"] = "Project management"
            entry["Supplier"] = "Transportation"
            entry["Description"] = "Flight"
            filled = True

        if "TONY S" in desc or "Tony S" in desc or "tony s" in desc or "eat" in desc or "EAT" in desc:
            entry["Type"] = "General"
            entry["Supplier"] = "F&B"
            entry["Description"] = "F&B"
            filled = True

        if any(word in desc for word in ["AEGEANWEB", "AEGEAN", "OLYMPIC", "SKY", "ISRAIR", "WIZZ"]):
            entry["Type"] = "Project management"
            entry["Supplier"] = "Transportation"
            entry["Description"] = "Flight"
            filled = True

        if any(word in desc for word in ["DINNER", "FOOD", "CAFE", "COFFEE", "LUNCH", "BREAKFAST"]):
            entry["Type"] = "General"
            entry["Supplier"] = "F&B"
            entry["Description"] = "F&B"
            filled = True

        if ("broker" in desc or "Broker" in desc or "BROKER" in desc) and ("villa 1" in desc or "Villa 1" in desc or "VILLA 1" in desc):
            entry["Type"] = "Brokers"
            entry["Supplier"] = "Buyer Villa 1"
            entry["Description"] = "Broker fees"
            filled = True

        if ("broker" in desc or "Broker" in desc or "BROKER" in desc) and ("villa 2" in desc or "Villa 2" in desc or "VILLA 2" in desc):
            entry["Type"] = "Brokers"
            entry["Supplier"] = "Buyer Villa 2"
            entry["Description"] = "Broker fees"
            filled = True

        if "RF919086180000334" in desc:
            entry["Plot"] = "R4"
            entry["Expenses Type"] = "Soft Cost"
            entry["Type"] = "Utility Bills"
            entry["Supplier"] = "Municipality"
            entry["Description"] = "Electricity"
            filled = True

        if ("broker" in desc or "Broker" in desc or "BROKER" in desc) and ("villa 3" in desc or "Villa 3" in desc or "VILLA 3" in desc):
            entry["Type"] = "Brokers"
            entry["Supplier"] = "Buyer Villa 3"
            entry["Description"] = "Broker fees"
            filled = True

        if ("broker" in desc or "Broker" in desc or "BROKER" in desc) and ("villa 4" in desc or "Villa 4" in desc or "VILLA 4" in desc):
            entry["Type"] = "Brokers"
            entry["Supplier"] = "Buyer Villa 4"
            entry["Description"] = "Broker fees"
            filled = True

        if ("broker" in desc or "Broker" in desc or "BROKER" in desc) and ("villa 5" in desc or "Villa 5" in desc or "VILLA 5" in desc):
            entry["Type"] = "Brokers"
            entry["Supplier"] = "Buyer Villa 5"
            entry["Description"] = "Broker fees"
            filled = True

        if ("broker" in desc or "Broker" in desc or "BROKER" in desc) and ("villa 6" in desc or "Villa 6" in desc or "VILLA 6" in desc):
            entry["Type"] = "Brokers"
            entry["Supplier"] = "Buyer Villa 6"
            entry["Description"] = "Broker fees"
            filled = True

        if "GOOGLE" in desc or "ΣΥΝΔΡΟΜΗ ADVANCED FOR BUSINES" in desc:
            entry["Type"] = "Marketing"
            entry["Supplier"] = "Marketing"
            entry["Description"] = "Marketing Services fee"
            filled = True

        if "CRM" in desc:
            entry["Type"] = "Marketing"
            entry["Supplier"] = "reWire"
            entry["Description"] = "CRM"
            filled = True

        if "RF91908618000033404472101" in desc or "PROT-RF549086180000334" in desc:
            entry["Type"] = "Utility Bills"
            entry["Supplier"] = "Municipality"
            entry["Description"] = "Electricity"
            entry["Plot"] = "G2"
            filled = True

        if "RF38908618000033404445701" in desc or "RF389086180000334044" in desc:
            entry["Type"] = "Utility Bills"
            entry["Supplier"] = "Municipality"
            entry["Description"] = "Electricity"
            entry["Plot"] = "Y3"
            filled = True

        if "RF91908618000033404472101" in desc or "PROT-919086180000334" in desc:
            entry["Type"] = "Utility Bills"
            entry["Supplier"] = "Municipality"
            entry["Description"] = "Electricity"
            entry["Plot"] = "R4"
            filled = True

        if "UBER" in desc or "TAXI" in desc:
            entry["Type"] = "Project management"
            entry["Supplier"] = "Transportation"
            entry["Description"] = "Athens Taxi"
            filled = True

        if "OPENAI" in desc:
            entry["Type"] = "General"
            entry["Supplier"] = "Office expenses"
            entry["Description"] = "Office expense"
            filled = True

        if "TAG" in desc:
            entry["Type"] = "Architect"
            entry["Supplier"] = "TAG ARCHITECTS"
            if "SUP" in desc:
                entry["Description"] = "Supervision"
            else:
                entry["Description"] = "Planning"
            filled = True

        if "OASA" in desc:
            entry["Type"] = "Project management"
            entry["Supplier"] = "Transportation"
            entry["Description"] = "Transportation"
            filled = True

        if "ΔΗΜΟ-RF369029090000097" in desc:
            entry["Type"] = "Utility Bills"
            entry["Supplier"] = "Municipality"
            entry["Description"] = "Water"
            entry["Plot"] = "Y3"
            filled = True

        if any(term in desc for term in ["MANAGEMENT", "MANAG.", "MGMT", "MNGMT"]) and row['ΠΟΣΟ'] in [-1550, -1550.00, -1550.0, 1550.00, 1550.0, 2055, 2055.0, 2057.0, 1550]:
            entry["Type"] = "Worker 1"
            entry["Supplier"] = "Aiolos Athens"
            entry["Description"] = "management fees"
            filled = True

        if any(term in desc for term in ["COSM", "COSMOTE", "PHONE"]):
            entry["Type"] = "Hotel operation"
            entry["Supplier"] = "Cosmote"
            entry["Description"] = "Telephone"
            filled = True

        if "RF389086180000334" in desc:
            entry["Type"] = "Utility Bills"
            entry["Supplier"] = "Municipality"
            entry["Description"] = "Electricity"
            entry["Plot"] = "Y3"
            filled = True

        # ============================================
        # END OF Ilisia RULES
        # ============================================

        if not filled:
            entry["Description"] = f"🟨 {entry['Description']}"

        results.append(entry)

    return pd.DataFrame(results)


# ============================================
# MAIN TABS
# ============================================
tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs([
    "📊 Excel Classifier", 
    "📝 Payment Instructions", 
    "🧾 Invoices",
    "📋 Receipts Database",
    "📅 Delay Penalties",
    "ℹ️ Help"
])

# ============================================
# TAB 1: EXCEL CLASSIFIER
# ============================================
with tab1:
    col1, col2 = st.columns([2, 1])
    
    with col1:
        st.markdown('<div class="info-card">', unsafe_allow_html=True)
        st.markdown("### 📊 Excel Classifier")
        st.markdown("Upload your financial Excel file to automatically categorize and organize transactions.")
        
       
        format_type = st.selectbox(
            "Select File Format",
            ["Diakofti", "Athens", "Ilisia NBG"],
            help="Choose the format that matches your data structure"
        )
        
        # File uploader
        uploaded_file = st.file_uploader(
            "Upload Excel or CSV File",
            type=["xlsx", "csv", "xls"],
            help="Drag and drop or click to browse"
        )
        st.markdown('</div>', unsafe_allow_html=True)
    
    with col2:
        st.markdown('<div class="info-card">', unsafe_allow_html=True)
        st.markdown("### 📌 Quick Guide")
        st.markdown("""
        **Formats:**
        - **Diakofti**: Plot-based transactions
        - **Athens**: Office transactions
        - **Ilisia NBG**: Ilisia project transactions
        
        **Supported Files:**
        - Excel (.xlsx, .xls)
        - CSV files
        
        **Output:**
        - Auto-categorized data
        - Entries needing review marked with 🟨
        """)
        st.markdown('</div>', unsafe_allow_html=True)
    
    # Process uploaded file
    if uploaded_file:
        st.markdown('<div class="success-msg">✅ File uploaded successfully!</div>', unsafe_allow_html=True)
        
        # Process button
        if st.button("🚀 Process File", use_container_width=True, key="process_excel"):
            with st.spinner("Processing your data..."):
                try:
                    # Read file
                    if uploaded_file.name.lower().endswith('.csv'):
                        if format_type == "Diakofti":
                            df = pd.read_csv(uploaded_file, encoding="ISO-8859-7")
                        else:
                            # נסה UTF-8 ואז נפילה ל-ISO-8859-7 במקרה של יוונית
                            try:
                                df = pd.read_csv(uploaded_file)
                            except UnicodeDecodeError:
                                df = pd.read_csv(uploaded_file, encoding="ISO-8859-7")
                    else:
                        df = pd.read_excel(uploaded_file)
                    
                    # Process based on format
                    if format_type == "Diakofti":
                        result_df = process_diakofti_file(df)
                    elif format_type == "Athens":
                        result_df = process_athens_file(df)
                    else:  # Ilisia NBG
                        result_df = process_ilisia_file(df)
                    
                    # Calculate metrics
                    total_entries = len(result_df)
                    needs_review = result_df['Description'].astype(str).str.contains('🟨').sum() if total_entries else 0
                    auto_classified = total_entries - needs_review
                    success_rate = (auto_classified / total_entries * 100) if total_entries > 0 else 0
                    
                    # Display metrics
                    st.markdown("""
                    <div class="metric-container">
                        <div class="metric-box">
                            <div class="metric-value">{}</div>
                            <div class="metric-label">Total Entries</div>
                        </div>
                        <div class="metric-box">
                            <div class="metric-value">{}</div>
                            <div class="metric-label">Auto-Classified</div>
                        </div>
                        <div class="metric-box">
                            <div class="metric-value">{}</div>
                            <div class="metric-label">Need Review</div>
                        </div>
                        <div class="metric-box">
                            <div class="metric-value">{:.1f}%</div>
                            <div class="metric-label">Success Rate</div>
                        </div>
                    </div>
                    """.format(total_entries, auto_classified, needs_review, success_rate), unsafe_allow_html=True)
                    
                    # Show preview
                    st.markdown("### 📋 Data Preview")
                    st.dataframe(result_df.head(10), use_container_width=True)
                    
                    # Download button
                    output = BytesIO()
                    result_df.to_excel(output, index=False, engine='openpyxl')
                    output.seek(0)
                    
                    st.download_button(
                        label="📥 Download Processed File",
                        data=output,
                        file_name=f"{format_type.lower()}_processed_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        use_container_width=True
                    )
                    
                except Exception as e:
                    st.error(f"❌ Error processing file: {str(e)}")

# ============================================
# TAB 2: PAYMENT INSTRUCTIONS
# ============================================
with tab2:
    col1, col2 = st.columns([2, 1])
    
    with col1:
        st.markdown('<div class="info-card">', unsafe_allow_html=True)
        st.markdown("### 📝 Generate Payment Instructions")
        
        # Project selection
        project = st.selectbox(
            "Select Project",
            sorted(set(k[0] for k in VILLA_OWNERS)),
            help="Choose the project plot",
            key="payment_project"
        )
        
        # Villa selection
        villa_options = sorted(set(k[1] for k in VILLA_OWNERS if k[0] == project))
        villa = st.selectbox(
            "Select Villa",
            villa_options,
            help="Select the villa number",
            key="payment_villa"
        )
        
        # Display owner name
        client_name = VILLA_OWNERS.get((project, villa), "")
        if client_name:
            st.markdown(f'<div class="info-msg">👤 <strong>Owner:</strong> {client_name}</div>', unsafe_allow_html=True)
        
        # Payment details
        st.markdown("### 💳 Payment Details")
        payment_order = st.text_input("Payment Order Number", placeholder="e.g., 12345", key="payment_order")
        amount = st.text_input("Amount in Euro (€)", placeholder="e.g., 5000", key="payment_amount")
        extra_text = st.text_area("Additional Notes (Optional)", placeholder="Any additional payment information...", key="payment_notes")
        
        st.markdown('</div>', unsafe_allow_html=True)
    
    with col2:
        st.markdown('<div class="template-preview">', unsafe_allow_html=True)
        st.markdown("### 📄 Template Preview")
        st.markdown("See how your payment instruction will look:")
        # Show template preview image from GitHub
        st.image("https://raw.githubusercontent.com/Daniel43450/aiolos-excel-app/main/payment_order_PNG.png", 
                 caption="Payment Instruction Template",
                 use_container_width=True)
        st.markdown('</div>', unsafe_allow_html=True)
    
    # Generate payment instruction
    if st.button("📄 Generate Payment Instruction", use_container_width=True, key="generate_payment"):
        if payment_order and amount:
            try:
                # Note: You'll need to have the template file
                template = Document("default_tempate.docx")
                
                # Replace placeholders
                for p in template.paragraphs:
                    p.text = p.text.replace("{{date}}", datetime.datetime.now().strftime("%d/%m/%Y"))
                    p.text = p.text.replace("{{plot}}", project)
                    p.text = p.text.replace("{{villa_no}}", villa)
                    p.text = p.text.replace("{{client_name}}", client_name)
                    p.text = p.text.replace("{{payment_order_number}}", payment_order)
                    p.text = p.text.replace("{{sum}}", amount)
                    p.text = p.text.replace("{{Extra Payment text}}", extra_text)
                
                # Save to buffer
                buffer = BytesIO()
                template.save(buffer)
                buffer.seek(0)
                
                filename = f"Payment_Instruction_{project}_{villa}_Order_{payment_order}.docx"
                
                # שמור את ההוראה בהיסטוריה
                payment_instruction = {
                    "id": f"PI_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}",
                    "project": project,
                    "villa": villa,
                    "client_name": client_name,
                    "payment_order": payment_order,
                    "amount": amount,
                    "notes": extra_text,
                    "created_date": datetime.datetime.now().strftime("%d/%m/%Y %H:%M"),
                    "timestamp": datetime.datetime.now().isoformat()
                }
                
                st.session_state.payment_instructions_db.append(payment_instruction)
                
                # עדיין שמור גם כ-last_payment לתאימות לאחור
                st.session_state.last_payment = payment_instruction
                # Persist to disk so history survives refresh/close
                try:
                    with open("payment_instructions_db.json", "w", encoding="utf-8") as f:
                        json.dump(st.session_state.payment_instructions_db, f, ensure_ascii=False, indent=2)
                except Exception as e:
                    st.warning(f"Could not persist payment instructions: {e}")

        
                st.markdown('<div class="success-msg">✅ Payment instruction generated successfully!</div>', unsafe_allow_html=True)
                
                st.download_button(
                    label="📥 Download Payment Instruction",
                    data=buffer,
                    file_name=filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
                
                # Option to create invoice
                st.markdown("---")
                st.markdown("### 🧾 Next Step: Create Invoice")
                st.info("Payment instruction created! You can now create an invoice for this payment in the Invoices tab.")
                
            except FileNotFoundError:
                st.error("❌ Template file 'default_template.docx' not found. Please add it to the app directory.")
            except Exception as e:
                st.error(f"❌ Error generating payment instruction: {str(e)}")
        else:
            st.warning("⚠️ Please fill in Payment Order Number and Amount")
            
# ============================================
# TAB 3: INVOICES
# ============================================

with tab3:
    st.markdown('<div class="info-card">', unsafe_allow_html=True)
    st.markdown("### 🧾 Receipt of Funds Generator")
    
    # ---------- Persistence helpers for Payment Instructions ----------
    PI_DB_FILE = "payment_instructions_db.json"

    def _load_payment_db():
        try:
            if os.path.exists(PI_DB_FILE):
                with open(PI_DB_FILE, "r", encoding="utf-8") as f:
                    data = json.load(f)
                    return data if isinstance(data, list) else []
        except Exception:
            pass
        return []

    def _save_payment_db(data):
        try:
            with open(PI_DB_FILE, "w", encoding="utf-8") as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
        except Exception as e:
            st.warning(f"Could not persist payment instructions: {e}")

    # Bootstrap session state from disk or sync to disk
    if 'payment_instructions_db' not in st.session_state:
        st.session_state.payment_instructions_db = _load_payment_db()
    else:
        # אם יש נתונים בזיכרון נטען אותם לדיסק כדי להבטיח התמדה
        if st.session_state.payment_instructions_db:
            _save_payment_db(st.session_state.payment_instructions_db)
        else:
            # אם אין בזיכרון ננסה לטעון מהדיסק
            disk_data = _load_payment_db()
            if disk_data:
                st.session_state.payment_instructions_db = disk_data

    # ---------- Helpers for form state ----------
    def _all_projects():
        return sorted(set(k[0] for k in VILLA_OWNERS))

    def _villas_for(project):
        return sorted(set(k[1] for k in VILLA_OWNERS if k[0] == project))

    def _ensure_form_state_defaults():
        # init project
        if 'receipt_project' not in st.session_state:
            all_proj = _all_projects()
            st.session_state.receipt_project = all_proj[0] if all_proj else ""

        # init villa
        if 'receipt_villa' not in st.session_state:
            villas = _villas_for(st.session_state.receipt_project)
            st.session_state.receipt_villa = villas[0] if villas else ""

        # init text fields
        if 'receipt_payment_order' not in st.session_state:
            st.session_state.receipt_payment_order = ""
        if 'receipt_amount' not in st.session_state:
            st.session_state.receipt_amount = ""
        if 'receipt_extra_text' not in st.session_state:
            st.session_state.receipt_extra_text = ""

        # init date
        if 'receipt_date' not in st.session_state or not isinstance(st.session_state.receipt_date, datetime.date):
            st.session_state.receipt_date = datetime.date.today()

    def _load_into_form_from_instruction(instr):
        # Fill widget state directly from selected instruction
        st.session_state.receipt_project = instr.get('project', _all_projects()[0]) or _all_projects()[0]
        villas = _villas_for(st.session_state.receipt_project)
        v = instr.get('villa') if instr.get('villa') in villas else (villas[0] if villas else "")
        st.session_state.receipt_villa = v
        st.session_state.receipt_payment_order = str(instr.get('payment_order', ''))
        st.session_state.receipt_amount = str(instr.get('amount', ''))
        st.session_state.receipt_date = datetime.date.today()  # תמיד להיום
        st.session_state.receipt_extra_text = instr.get('notes', '')

    def _clear_form():
        # נקה את כל מפתחות הטופס כדי שהאתחול ימלא ערכים חוקיים מחדש
        for k in [
            "receipt_project",
            "receipt_villa",
            "receipt_payment_order",
            "receipt_amount",
            "receipt_date",
            "receipt_extra_text",
            "load_into_form",
        ]:
            if k in st.session_state:
                st.session_state.pop(k)
        if 'selected_payment_instruction' in st.session_state:
            st.session_state.pop('selected_payment_instruction')

    def _find_template_path():
        for pth in ["Receipt_of_Funds.docx", "/mnt/data/Receipt_of_Funds.docx"]:
            if os.path.exists(pth):
                return pth
        return None

    def _fill_template_docx(template_path, mapping):
        from docx import Document
        doc = Document(template_path)

        # paragraphs
        for p in doc.paragraphs:
            txt = p.text
            for k, v in mapping.items():
                txt = txt.replace(k, v)
            p.text = txt  # resets runs

        # tables
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        txt = p.text
                        for k, v in mapping.items():
                            txt = txt.replace(k, v)
                        p.text = txt

        # keep title bold
        for p in doc.paragraphs:
            if p.text.strip().startswith("Receipt of Funds"):
                for run in p.runs:
                    run.bold = True
        return doc

    # ---------- Payment Instructions History ----------
    if st.session_state.payment_instructions_db:
        st.markdown("### 📋 Payment Instructions History")
        st.markdown('<div class="info-msg">💡 Select a payment instruction to auto-fill the receipt form</div>', unsafe_allow_html=True)
        
        cols_history = st.columns([0.5, 1, 1.5, 1.5, 1, 1, 0.5, 0.5])
        with cols_history[0]: st.write("**#**")
        with cols_history[1]: st.write("**Date**")
        with cols_history[2]: st.write("**Project**")
        with cols_history[3]: st.write("**Client**")
        with cols_history[4]: st.write("**Order#**")
        with cols_history[5]: st.write("**Amount**")
        with cols_history[6]: st.write("**Load**")
        with cols_history[7]: st.write("**Delete**")
        st.markdown("---")
        
        for idx, instruction in enumerate(reversed(st.session_state.payment_instructions_db)):
            actual_idx = len(st.session_state.payment_instructions_db) - 1 - idx
            cols_row = st.columns([0.5, 1, 1.5, 1.5, 1, 1, 0.5, 0.5])
            with cols_row[0]:
                st.write(f"{idx + 1}")
            with cols_row[1]:
                st.write(instruction['created_date'].split(' ')[0])
            with cols_row[2]:
                st.write(f"{instruction['project']} - {instruction['villa']}")
            with cols_row[3]:
                client_display = instruction['client_name'][:25] + "..." if len(instruction['client_name']) > 25 else instruction['client_name']
                st.write(client_display)
            with cols_row[4]:
                st.write(instruction['payment_order'])
            with cols_row[5]:
                st.write(f"€{instruction['amount']}")
            with cols_row[6]:
                if st.button("📥", key=f"load_pi_{actual_idx}", help="Load this payment instruction"):
                    st.session_state.selected_payment_instruction = instruction
                    st.session_state.load_into_form = True   # flag to prefill widgets
                    st.rerun()
            with cols_row[7]:
                if st.button("🗑️", key=f"delete_pi_{actual_idx}", help="Delete this payment instruction"):
                    st.session_state.payment_instructions_db.pop(actual_idx)
                    _save_payment_db(st.session_state.payment_instructions_db)  # persist deletion
                    st.success("Payment instruction deleted!")
                    st.rerun()
        st.markdown("---")

    # ---------- Defaults & auto-fill after load ----------
    _ensure_form_state_defaults()

    selected_instruction = st.session_state.get('selected_payment_instruction')
    if st.session_state.get('load_into_form') and selected_instruction:
        _load_into_form_from_instruction(selected_instruction)
        st.session_state.load_into_form = False
        st.markdown('<div class="success-msg">✅ Payment instruction loaded! Fields filled automatically.</div>', unsafe_allow_html=True)
    elif 'last_payment' in st.session_state and not any(st.session_state.get(k) for k in ['selected_payment_instruction','load_into_form']):
        st.markdown('<div class="info-msg">📌 Recent payment instruction detected! You can load it from history above.</div>', unsafe_allow_html=True)

    # ---------- Form UI ----------
    col1, col2 = st.columns([3, 2])
    with col1:
        st.markdown("### 📝 Receipt of Funds Details")

        st.selectbox(
            "Project",
            _all_projects(),
            key="receipt_project"
        )

        villa_options_receipt = _villas_for(st.session_state.receipt_project)
        if not villa_options_receipt:
            st.session_state.receipt_villa = ""
            villa_options_display = [""]
        else:
            if st.session_state.receipt_villa not in villa_options_receipt:
                st.session_state.receipt_villa = villa_options_receipt[0]
            villa_options_display = villa_options_receipt

        st.selectbox(
            "Villa",
            villa_options_display,
            key="receipt_villa"
        )

        receipt_client = VILLA_OWNERS.get((st.session_state.receipt_project, st.session_state.receipt_villa), "")
        if receipt_client:
            st.markdown(f'<div class="info-msg">👤 <strong>Client:</strong> {receipt_client}</div>', unsafe_allow_html=True)

        st.text_input("Payment Order Number", placeholder="001", key="receipt_payment_order")
        st.text_input("Amount (€)", placeholder="5000", key="receipt_amount")
        st.date_input("Date of Receipt", key="receipt_date")
        st.text_area(
            "Extra Receipt Text (Optional)",
            placeholder="Additional information about the payment...",
            help="This text will appear in the receipt. Leave empty if not needed.",
            key="receipt_extra_text"
        )

        # Clear form really clears everything
        if st.button("🔄 Clear Form", use_container_width=True, key="clear_receipt_form"):
            _clear_form()
            st.rerun()

    with col2:
        st.markdown("### 📊 Receipt Preview")
        if st.session_state.receipt_payment_order and st.session_state.receipt_amount:
            st.markdown(f"""
            <div style="background: white; padding: 1.5rem; border-radius: 12px; border: 2px solid #e0e0e0; text-align: left;">
                <h4 style="color: #1e3c72; margin-bottom: 1rem;">🧾 Receipt of Funds Preview</h4>
                <p><strong>To:</strong> {receipt_client}</p>
                <p><strong>Project:</strong> {st.session_state.receipt_project}</p>
                <p><strong>Villa #:</strong> {st.session_state.receipt_villa}</p>
                <p><strong>Payment Order:</strong> {st.session_state.receipt_payment_order}</p>
                <p><strong>Date:</strong> {st.session_state.receipt_date.strftime('%d/%m/%Y')}</p>
                <p><strong>Amount:</strong> <span style="font-weight: bold;">€{st.session_state.receipt_amount}</span></p>
                {f'<p><strong>Extra Text:</strong> {st.session_state.receipt_extra_text[:50]}{"..." if len(st.session_state.receipt_extra_text) > 50 else ""}</p>' if st.session_state.receipt_extra_text else ""}
            </div>
            """, unsafe_allow_html=True)
        else:
            st.markdown("""
            <div style="background: #f8f9fa; padding: 1.5rem; border-radius: 12px; border: 2px dashed #dee2e6; text-align: center;">
                <h4 style="color: #6c757d;">📄 Receipt Preview</h4>
                <p style="color: #6c757d;">Fill in the details to see preview</p>
            </div>
            """, unsafe_allow_html=True)

    # ---------- Generate DOCX from template ----------
    if st.button("📥 Generate and Download Receipt of Funds", use_container_width=True, key="generate_receipt"):
        rp = st.session_state.receipt_project
        rv = st.session_state.receipt_villa
        rc = VILLA_OWNERS.get((rp, rv), "")
        rpo = st.session_state.receipt_payment_order
        ra = st.session_state.receipt_amount
        rd = st.session_state.receipt_date
        rx = st.session_state.receipt_extra_text

        if rpo and ra:
            try:
                template_path = _find_template_path()
                if not template_path:
                    raise FileNotFoundError("Template file 'Receipt_of_Funds.docx' not found. Place it in the app folder or /mnt/data.")
                
                mapping = {
                    "{{client_name}}": rc or "",
                    "{{plot}}": rp or "",
                    "{{villa_no}}": rv or "",
                    "{{payment_order_number}}": rpo or "",
                    "{{sum}}": str(ra or ""),
                    "{{date}}": rd.strftime("%d/%m/%Y") if rd else "",
                    "{{Extra Receipt text}}": (rx if rx and rx.strip() else "")
                }

                doc = _fill_template_docx(template_path, mapping)
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)

                # filename with two spaces before client_name
                filename = f"{rp} - Receipt of Funds {rpo} -  {rc}.docx"

                # Save a record so TAB 4 shows it
                receipt_record = {
                    "type": "Receipt of Funds",
                    "number": rpo,
                    "date": rd.strftime("%Y-%m-%d"),
                    "project": rp,
                    "villa": rv,
                    "client": rc,
                    "amount": ra,
                    "notes": rx,
                    "timestamp": datetime.datetime.now().isoformat()
                }
                st.session_state.receipts_db.append(receipt_record)

                st.markdown('<div class="success-msg">✅ Receipt of Funds generated successfully!</div>', unsafe_allow_html=True)
                st.download_button(
                    label="📥 Download Receipt of Funds (DOCX)",
                    data=buffer,
                    file_name=filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )

            except FileNotFoundError as e:
                st.error(f"❌ {str(e)}")
            except Exception as e:
                st.error(f"❌ Error generating receipt: {str(e)}")
        else:
            st.warning("⚠️ Please fill in Payment Order Number and Amount")
    
    st.markdown('</div>', unsafe_allow_html=True)

# ============================================
# TAB 4: RECEIPTS DATABASE
# ============================================
with tab4:
    st.markdown('<div class="info-card">', unsafe_allow_html=True)
    st.markdown("### 📋 All Receipts & Invoices Database")

    # ---------- Persistence helpers ----------
    RECEIPTS_DB_FILE = "receipts_db.json"

    def _load_receipts_db():
        try:
            if os.path.exists(RECEIPTS_DB_FILE):
                with open(RECEIPTS_DB_FILE, "r", encoding="utf-8") as f:
                    data = json.load(f)
                    return data if isinstance(data, list) else []
        except Exception:
            pass
        return []

    def _save_receipts_db(data):
        try:
            with open(RECEIPTS_DB_FILE, "w", encoding="utf-8") as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
        except Exception as e:
            st.warning(f"Could not persist receipts DB: {e}")

    # Bootstrap session state from disk or sync to disk
    if 'receipts_db' not in st.session_state:
        st.session_state.receipts_db = _load_receipts_db()
    else:
        if st.session_state.receipts_db:
            _save_receipts_db(st.session_state.receipts_db)
        else:
            disk_data = _load_receipts_db()
            if disk_data:
                st.session_state.receipts_db = disk_data

    # Combine all receipts and invoices
    all_records = st.session_state.receipts_db

    if all_records:
        # Create DataFrame for display
        df_records = pd.DataFrame(all_records)

        # Summary metrics
        total_records = len(df_records)
        total_amount = 0.0
        for r in all_records:
            try:
                amt = float(str(r.get('amount', '')).replace(',', '').strip())
                total_amount += amt
            except Exception:
                pass
        unique_projects = len(set(r.get('project', '') for r in all_records if r.get('project')))

        st.markdown(f"""
        <div class="metric-container">
            <div class="metric-box">
                <div class="metric-value">{total_records}</div>
                <div class="metric-label">Total Records</div>
            </div>
            <div class="metric-box">
                <div class="metric-value">€{total_amount:,.2f}</div>
                <div class="metric-label">Total Amount</div>
            </div>
            <div class="metric-box">
                <div class="metric-value">{unique_projects}</div>
                <div class="metric-label">Projects</div>
            </div>
        </div>
        """, unsafe_allow_html=True)
        
        # Filter options
        col1, col2, col3 = st.columns(3)
        with col1:
            filter_project = st.selectbox(
                "Filter by Project",
                ["All"] + sorted(set(r.get('project', '') for r in all_records if r.get('project'))),
                key="filter_project"
            )
        with col2:
            filter_villa = st.selectbox(
                "Filter by Villa",
                ["All"] + sorted(set(r.get('villa', '') for r in all_records if r.get('villa'))),
                key="filter_villa"
            )
        with col3:
            filter_type = st.selectbox(
                "Filter by Type",
                ["All", "Invoice", "Receipt of Funds", "Payment Instruction"],
                key="filter_type"
            )
        
        # Apply filters
        filtered_records = all_records
        if filter_project != "All":
            filtered_records = [r for r in filtered_records if r.get('project') == filter_project]
        if filter_villa != "All":
            filtered_records = [r for r in filtered_records if r.get('villa') == filter_villa]
        if filter_type != "All":
            filtered_records = [r for r in filtered_records if r.get('type') == filter_type]
        
        # Display table
        if filtered_records:
            st.markdown("### 📊 Records Table")

            for idx, record in enumerate(filtered_records):
                display_type = record.get('type', '')
                display_number = record.get('number') or record.get('payment_order') or record.get('invoice_number') or "N/A"
                display_project = record.get('project', '')
                display_villa = record.get('villa', '')
                display_amount = record.get('amount', '')
                with st.expander(f"{display_type} #{display_number} - {display_project} {display_villa} - €{display_amount}"):
                    colA, colB = st.columns([3, 1])
                    with colA:
                        st.write(f"**Date:** {record.get('date', '')}")
                        st.write(f"**Client:** {record.get('client', '')}")
                        st.write(f"**Amount:** €{display_amount}")
                        notes_val = record.get('notes') or record.get('extra_text') or ""
                        if notes_val:
                            st.write(f"**Notes:** {notes_val}")
                    with colB:
                        # מחיקה בטוחה + התמדה לדיסק
                        if st.button("🗑️ Delete", key=f"delete_{idx}"):
                            for i, r0 in enumerate(st.session_state.receipts_db):
                                if r0 is record:
                                    st.session_state.receipts_db.pop(i)
                                    _save_receipts_db(st.session_state.receipts_db)
                                    break
                            st.rerun()
        else:
            st.info("No records found with the selected filters.")
        
        # Export options
        st.markdown("---")
        st.markdown("### 📥 Export Options")
        
        if st.button("📊 Export All to Excel", use_container_width=True):
            df_export = pd.DataFrame(all_records)
            output = BytesIO()
            df_export.to_excel(output, index=False, engine='openpyxl')
            output.seek(0)
            
            st.download_button(
                label="📥 Download Excel File",
                data=output,
                file_name=f"receipts_database_{datetime.datetime.now().strftime('%Y%m%d')}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True
            )
    else:
        st.info("📭 No records found yet. Create your first one in the Payment Instructions or Invoices tab!")
    
    st.markdown('</div>', unsafe_allow_html=True)
    
    # Clear database option
    st.markdown("---")
    st.markdown("### ⚠️ Database Management")
    col1, col2, col3 = st.columns([1, 1, 2])
    with col1:
        if st.button("🗑️ Clear All Records", use_container_width=True):
            st.session_state.show_clear_confirm = True
    
    if 'show_clear_confirm' in st.session_state and st.session_state.show_clear_confirm:
        with col2:
            if st.button("✅ Confirm Clear", use_container_width=True):
                st.session_state.receipts_db = []
                _save_receipts_db(st.session_state.receipts_db)  # persist empty DB
                # אופציונלי: נקה גם invoices_db אם תרצה התמדה נפרדת לקבצים
                st.session_state.invoices_db = []
                st.session_state.show_clear_confirm = False
                st.success("Database cleared successfully!")
                st.rerun()
        with col3:
            if st.button("❌ Cancel", use_container_width=True):
                st.session_state.show_clear_confirm = False
                st.rerun()


# ============================================
# TAB 5: DELAY PENALTIES CALCULATOR
# ============================================

with tab5:
    st.markdown('<div class="info-card">', unsafe_allow_html=True)
    st.markdown("### ⏱️ Delay Penalties Calculator")

    # ---------- Imports for PDF ----------
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import mm
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
        REPORTLAB_OK = True
    except Exception as _e:
        REPORTLAB_OK = False
        reportlab_err = str(_e)

    # ---------- Helpers ----------
    MONTH_NAMES = {
        1: "January", 2: "February", 3: "March", 4: "April",
        5: "May", 6: "June", 7: "July", 8: "August",
        9: "September", 10: "October", 11: "November", 12: "December"
    }
    DAY_NAMES = ["Sunday", "Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday"]

    def _default_monthly_rates():
        return {
            1: 75.0, 2: 75.0, 3: 75.0, 4: 75.0, 5: 75.0,
            6: 175.0, 7: 250.0, 8: 250.0, 9: 175.0, 10: 75.0, 11: 75.0, 12: 75.0
        }

    def _compute_penalties(original_date, actual_date, grace_days, villas, monthly_rates, exclude_saturday, currency_symbol):
        rows = []
        cur = original_date
        grace_end = original_date + datetime.timedelta(days=grace_days)
        cumulative = 0.0

        while cur <= actual_date:
            weekday_idx = cur.weekday()  # Mon=0..Sun=6
            weekday_for_table = (weekday_idx + 1) % 7  # Sun=0..Sat=6

            in_grace = cur <= grace_end
            is_saturday = weekday_for_table == 6
            month_idx = cur.month
            per_villa_rate = 0.0 if in_grace or (exclude_saturday and is_saturday) else float(monthly_rates.get(month_idx, 0.0))
            daily_total = per_villa_rate * villas
            if not in_grace and not (exclude_saturday and is_saturday):
                cumulative += daily_total

            if cur < original_date:
                status = "Before original date"
            elif in_grace:
                status = "Grace period"
            elif exclude_saturday and is_saturday:
                status = "Excluded (Saturday)"
            else:
                status = "Penalty period"

            rows.append({
                "Date": cur.strftime("%d/%m/%Y"),
                "Weekday": DAY_NAMES[weekday_for_table],
                "Month": MONTH_NAMES[month_idx],
                "Status": status,
                "Per-villa daily": f"{currency_symbol}{per_villa_rate:,.2f}",
                "Daily total": f"{currency_symbol}{daily_total:,.2f}",
                "Cumulative": f"{currency_symbol}{cumulative:,.2f}",
            })
            cur += datetime.timedelta(days=1)

        df = pd.DataFrame(rows)

        total_days_late = max(0, (actual_date - original_date).days)
        grace_len = max(0, (grace_end - original_date).days)
        penalty_days = max(0, total_days_late - grace_len)
        total_penalty_val = cumulative

        return df, {
            "original": original_date.strftime("%d/%m/%Y"),
            "actual": actual_date.strftime("%d/%m/%Y"),
            "days_late": total_days_late,
            "grace_days": grace_len,
            "penalty_days": penalty_days,
            "villas": villas,
            "total_penalty": f"{currency_symbol}{total_penalty_val:,.2f}",
        }

    def _to_csv_bytes(df, header_rows):
        from io import StringIO
        sio = StringIO()
        sio.write("\ufeff")
        for k, v in header_rows:
            sio.write(f"{k},{v}\n")
        sio.write("\n")
        df.to_csv(sio, index=False)
        return sio.getvalue().encode("utf-8")

    def _printable_html(df, summary, monthly_rates, currency_symbol):
        rates_html = "".join(
            f"<div style='background:#f8f9fa;padding:8px;border-radius:4px;text-align:center;border:1px solid #eee;'>"
            f"<strong>{MONTH_NAMES[m]}:</strong> {currency_symbol}{monthly_rates[m]:,.2f}</div>"
            for m in range(1, 13)
        )
        table_html = df.to_html(index=False, border=0, classes="dataframe", escape=False)
        html = f"""
<!DOCTYPE html>
<html lang="en" dir="ltr">
<head>
<meta charset="UTF-8">
<title>Delay Penalties Report</title>
<style>
body {{ font-family: Arial, sans-serif; margin: 20px; color: #222; }}
h1 {{ text-align: center; color: #1e3c72; border-bottom: 2px solid #2a5298; padding-bottom: 10px; }}
.section {{ background: #f8f9fa; padding: 14px; border-radius: 8px; border: 1px solid #eee; margin: 16px 0; }}
.grid-6 {{ display: grid; grid-template-columns: repeat(6, 1fr); gap: 10px; }}
.summary {{ background: #1e3c72; color: white; padding: 14px; border-radius: 8px; margin: 16px 0; text-align: center; font-weight: bold; }}
table.dataframe {{ width: 100%; border-collapse: collapse; font-size: 12px; }}
table.dataframe th, table.dataframe td {{ border: 1px solid #ddd; padding: 6px; text-align: center; }}
table.dataframe thead th {{ background: #2a5298; color: #fff; }}
@media print {{
  body {{ margin: 10mm; }}
  table.dataframe {{ font-size: 10px; }}
  table.dataframe th, table.dataframe td {{ padding: 4px; }}
}}
</style>
</head>
<body>
  <h1>Delay Penalties Report</h1>

  <div class="section">
    <h3>Calculation Settings</h3>
    <p><strong>Original delivery date:</strong> {summary["original"]}</p>
    <p><strong>Actual delivery date:</strong> {summary["actual"]}</p>
    <p><strong>Grace period (days):</strong> {summary["grace_days"]}</p>
    <p><strong>Number of villas:</strong> {summary["villas"]}</p>
    <p><strong>Currency:</strong> {currency_symbol}</p>
    <h4>Monthly penalty rates (per day per villa)</h4>
    <div class="grid-6">{rates_html}</div>
  </div>

  <div class="summary">
    Total late days: {summary["days_late"]} &nbsp;|&nbsp;
    Grace days: {summary["grace_days"]} &nbsp;|&nbsp;
    Penalty days: {summary["penalty_days"]} &nbsp;|&nbsp;
    <strong>Total penalty: {summary["total_penalty"]}</strong>
  </div>

  {table_html}

  <p style="text-align:center;color:#666;margin-top:20px;">Generated on: {datetime.datetime.now().strftime("%d/%m/%Y %H:%M")}</p>
</body>
</html>
"""
        return html

    def _to_pdf_bytes(df, summary, monthly_rates, currency_symbol):
        """Create a nicely formatted PDF with ReportLab."""
        if not REPORTLAB_OK:
            raise RuntimeError(f"reportlab is not available: {reportlab_err}")

        buf = BytesIO()
        doc = SimpleDocTemplate(
            buf,
            pagesize=A4,
            leftMargin=15*mm, rightMargin=15*mm,
            topMargin=15*mm, bottomMargin=15*mm
        )
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            "Title",
            parent=styles["Title"],
            textColor=colors.HexColor("#1e3c72"),
            fontSize=20,
            leading=24,
            alignment=1  # center
        )
        h3_style = ParagraphStyle("H3", parent=styles["Heading3"], textColor=colors.HexColor("#2a5298"))
        normal = styles["Normal"]

        story = []
        story.append(Paragraph("Delay Penalties Report", title_style))
        story.append(Spacer(1, 6))

        # Summary box
        summary_text = (
            f"<b>Original delivery date:</b> {summary['original']} &nbsp;&nbsp; "
            f"<b>Actual delivery date:</b> {summary['actual']} &nbsp;&nbsp; "
            f"<b>Grace days:</b> {summary['grace_days']} &nbsp;&nbsp; "
            f"<b>Penalty days:</b> {summary['penalty_days']} &nbsp;&nbsp; "
            f"<b>Villas:</b> {summary['villas']} &nbsp;&nbsp; "
            f"<b>Total penalty:</b> {summary['total_penalty']}"
        )
        story.append(Paragraph(summary_text, normal))
        story.append(Spacer(1, 10))

        # Monthly rates table
        story.append(Paragraph("Monthly penalty rates (per day per villa)", h3_style))
        rates_data = [["Month", "Rate"]]
        for m in range(1, 13):
            rates_data.append([MONTH_NAMES[m], f"{currency_symbol}{monthly_rates[m]:,.2f}"])
        rates_tbl = Table(rates_data, colWidths=[60*mm, 40*mm])
        rates_tbl.setStyle(TableStyle([
            ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#2a5298")),
            ("TEXTCOLOR", (0,0), (-1,0), colors.white),
            ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
            ("ALIGN", (0,0), (-1,0), "CENTER"),
            ("GRID", (0,0), (-1,-1), 0.25, colors.grey),
            ("ROWBACKGROUNDS", (0,1), (-1,-1), [colors.whitesmoke, colors.HexColor("#f8f9fa")]),
        ]))
        story.append(rates_tbl)
        story.append(Spacer(1, 10))

        # Daily breakdown table
        story.append(Paragraph("Daily breakdown", h3_style))

        # Convert df to raw data for table
        headers = list(df.columns)
        data = [headers] + df.values.tolist()

        # Limit rows to keep PDF size reasonable. Adjust if needed.
        MAX_ROWS = 1500
        if len(data) > MAX_ROWS + 1:
            data = data[:MAX_ROWS + 1]

        # Wider page table
        col_widths = [24*mm, 24*mm, 28*mm, 40*mm, 28*mm, 28*mm, 30*mm]
        tbl = Table(data, colWidths=col_widths, repeatRows=1)
        tbl_style = TableStyle([
            ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#2a5298")),
            ("TEXTCOLOR", (0,0), (-1,0), colors.white),
            ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
            ("ALIGN", (0,0), (-1,0), "CENTER"),
            ("GRID", (0,0), (-1,-1), 0.25, colors.grey),
            ("FONTSIZE", (0,0), (-1,-1), 8),
            ("VALIGN", (0,0), (-1,-1), "MIDDLE"),
        ])
        tbl.setStyle(tbl_style)
        story.append(tbl)

        doc.build(story)
        buf.seek(0)
        return buf.getvalue()

    # ---------- UI ----------
    c1, c2 = st.columns([2, 1])
    with c1:
        st.markdown("#### General Settings")
        left, right = st.columns(2)
        with left:
            original_date = st.date_input("Original delivery date", value=datetime.date(2025, 6, 1), key="pen_original_date")
            grace_days = st.number_input("Grace period (days)", min_value=0, value=61, step=1, key="pen_grace")
            villas = st.number_input("Number of villas", min_value=1, value=8, step=1, key="pen_villas")
        with right:
            actual_date = st.date_input("Actual delivery date", value=datetime.date(2025, 10, 1), key="pen_actual_date")
            currency_symbol = st.selectbox("Currency", ["€", "₪", "$"], index=0, key="pen_currency")
            exclude_saturday = st.checkbox("Exclude Saturdays from penalty", value=False, key="pen_ex_sat")

        st.markdown("#### Monthly penalty rates (per day per villa)")
        rates_grid1, rates_grid2, rates_grid3, rates_grid4, rates_grid5, rates_grid6 = st.columns(6)
        if "pen_rates" not in st.session_state:
            st.session_state.pen_rates = _default_monthly_rates()
        for idx, col in enumerate([rates_grid1, rates_grid2, rates_grid3, rates_grid4, rates_grid5, rates_grid6], start=1):
            with col:
                for m in range(idx, 13, 6):  # 1,7 ; 2,8 ; ...
                    st.session_state.pen_rates[m] = st.number_input(
                        f"{MONTH_NAMES[m]}",
                        min_value=0.0, step=1.0, value=float(st.session_state.pen_rates[m]),
                        key=f"pen_rate_{m}"
                    )

    with c2:
        st.markdown("#### Actions")
        calc_btn = st.button("🔄 Calculate penalties", use_container_width=True, key="pen_calc_btn")
        metrics_box = st.empty()

    st.markdown("---")

    # Run calculation on click or first render
    if calc_btn or "pen_last_df" not in st.session_state:
        try:
            df, summary = _compute_penalties(
                original_date,
                actual_date,
                grace_days,
                villas,
                st.session_state.pen_rates,
                exclude_saturday,
                currency_symbol
            )
            st.session_state.pen_last_df = df
            st.session_state.pen_last_summary = summary
        except Exception as e:
            st.error(f"Calculation error: {e}")

    if "pen_last_df" in st.session_state:
        df = st.session_state.pen_last_df
        summary = st.session_state.pen_last_summary

        # Metrics
        with c2:
            with metrics_box.container():
                st.markdown("""
                <div class="metric-container">
                    <div class="metric-box">
                        <div class="metric-value">{days}</div>
                        <div class="metric-label">Total Late Days</div>
                    </div>
                    <div class="metric-box">
                        <div class="metric-value">{grace}</div>
                        <div class="metric-label">Grace Days</div>
                    </div>
                    <div class="metric-box">
                        <div class="metric-value">{pen}</div>
                        <div class="metric-label">Penalty Days</div>
                    </div>
                </div>
                """.format(
                    days=summary["days_late"],
                    grace=summary["grace_days"],
                    pen=summary["penalty_days"]
                ), unsafe_allow_html=True)

        st.markdown(f"""
        <div class="info-msg">Original: <strong>{summary["original"]}</strong> &nbsp;|&nbsp;
        Actual: <strong>{summary["actual"]}</strong> &nbsp;|&nbsp;
        Villas: <strong>{summary["villas"]}</strong> &nbsp;|&nbsp;
        <span style="color:#1e3c72;">Total penalty: <strong>{summary["total_penalty"]}</strong></span></div>
        """, unsafe_allow_html=True)

        st.markdown("#### Daily breakdown")
        st.dataframe(df, use_container_width=True)

        # Downloads
        st.markdown("---")
        st.markdown("### 📥 Export")
        colA, colB, colC = st.columns(3)

        with colA:
            header_rows = [
                ("Original delivery date", summary["original"]),
                ("Actual delivery date", summary["actual"]),
                ("Grace period (days)", summary["grace_days"]),
                ("Penalty days", summary["penalty_days"]),
                ("Number of villas", summary["villas"]),
                ("Currency", currency_symbol),
                ("Total penalty", summary["total_penalty"]),
            ]
            csv_bytes = _to_csv_bytes(df, header_rows)
            st.download_button(
                label="📊 Download CSV",
                data=csv_bytes,
                file_name=f"delay_penalties_{actual_date.strftime('%Y%m%d')}.csv",
                mime="text/csv",
                use_container_width=True
            )

        with colB:
            html_str = _printable_html(df, summary, st.session_state.pen_rates, currency_symbol)
            html_bytes = html_str.encode("utf-8")
            st.download_button(
                label="🖨️ Download print view (HTML)",
                data=html_bytes,
                file_name=f"delay_penalties_{actual_date.strftime('%Y%m%d')}.html",
                mime="text/html",
                use_container_width=True
            )
            st.caption("Open the HTML and use your browser Print to save as PDF.")

        with colC:
            if REPORTLAB_OK:
                try:
                    pdf_bytes = _to_pdf_bytes(df, summary, st.session_state.pen_rates, currency_symbol)
                    st.download_button(
                        label="📄 Download PDF",
                        data=pdf_bytes,
                        file_name=f"delay_penalties_{actual_date.strftime('%Y%m%d')}.pdf",
                        mime="application/pdf",
                        use_container_width=True
                    )
                except Exception as e:
                    st.error(f"PDF generation error: {e}")
            else:
                st.warning("PDF export requires reportlab. Add to requirements: reportlab==3.6.13")

    st.markdown('</div>', unsafe_allow_html=True)


# ============================================
# TAB 6: HELP
# ============================================
with tab6:
    st.markdown('<div class="info-card">', unsafe_allow_html=True)
    st.markdown("### 🔧 How to Use Aiolos Financial Tools")
    
    st.markdown("""
    #### 📊 Excel Classifier
    1. **Select Format**: Choose between Diakofti (plot-based) or Athens (office) format
    2. **Upload File**: Upload your Excel or CSV file
    3. **Process**: Click the Process button to categorize transactions
    4. **Review**: Check entries marked with 🟨 - these need manual review
    5. **Download**: Download the processed file with all categorizations
    
    #### 📝 Payment Instructions
    1. **Select Villa**: Choose project and villa number
    2. **Enter Details**: Add payment order number and amount
    3. **Template Preview**: See how your document will look
    4. **Generate**: Click to create a Word document
    5. **Next Step**: Option to create an invoice for the payment
    
    #### 🧾 Invoices
    1. **Auto-fill**: If you created a payment instruction, details are pre-filled
    2. **Edit Details**: Adjust amount or add notes as needed
    3. **Generate**: Create invoice and save to database
    4. **Download**: Get a copy for your records
    
    #### 💰 Penalty Calculator
    1. **Set Dates**: Enter the original delivery date and actual delivery date  
    2. **Grace Period**: Define the number of grace days (no penalties)  
    3. **Penalty Rates**: Adjust penalty rates per villa for each month  
    4. **Calculate**: Click "Calculate Penalties" to generate the daily penalty table  
    5. **Review Table**: View status by day (Grace / Penalty) and cumulative fines  
    6. **Export**: Download results to Excel (CSV) or PDF report  

    #### 📋 Receipts Database
    1. **View All**: See all payment instructions and invoices
    2. **Filter**: Filter by project, villa, or type
    3. **Delete**: Remove individual records
    4. **Export**: Download entire database as Excel
    5. **Manage**: Clear database if needed
    
    #### 🔧 Adding Classification Rules
    To add classification rules, edit the code in the processing functions:
    - **Diakofti Rules**: Look for the section marked "🔴 Diakofti RULES"
    - **Athens Rules**: Look for the section marked "🔵 Athens RULES"
    
    #### 📞 Support
    For issues or questions, please contact the development team.
    
    ---
    
    ### 🏢 About Aiolos
    Aiolos Financial Tools is designed to streamline financial management for real estate projects,
    making it easy to process transactions, generate documents, and maintain records.
    """)
    
    st.markdown('</div>', unsafe_allow_html=True)
    
    # Version info
    st.markdown("---")
    col1, col2 = st.columns(2)
    with col1:
        st.markdown("**Version:** 3.0.0")
        st.markdown("**Last Updated:** " + datetime.datetime.now().strftime("%Y-%m-%d"))
    with col2:
        st.markdown("**Developed by:** Aiolos Team")
        st.markdown("**© 2024 Aiolos. All rights reserved.**")
